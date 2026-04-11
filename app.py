from flask import Flask, render_template, request, send_file, flash, redirect, url_for, Response, jsonify
import pandas as pd
import os
from werkzeug.utils import secure_filename
import tempfile
import io
import csv
from datetime import datetime
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import zipfile

# Importar módulo de banco de dados
from database import init_database, salvar_projeto, listar_projetos, carregar_resultados_projeto, excluir_projeto

app = Flask(__name__)
app.secret_key = 'sua_chave_secreta_aqui'

# Configurações para upload de arquivos
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'xlsx', 'csv'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Criar pasta de uploads se não existir
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Inicializar banco de dados
init_database()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    # Carregar projetos salvos na rota principal também
    projetos = listar_projetos()
    return render_template('dashboard_professional.html', projetos=projetos)

@app.route('/dashboard')
def dashboard():
    # Carregar projetos salvos
    projetos = listar_projetos()
    return render_template('dashboard_professional.html', projetos=projetos)

@app.route('/api/projetos')
def api_projetos():
    """API endpoint para buscar projetos salvos"""
    projetos = listar_projetos()
    return jsonify({'projetos': projetos})

@app.route('/tools')
def tools():
    return render_template('index.html')

@app.route('/tradicional')
def tradicional():
    return render_template('otimizador_cds.html')

@app.route('/otimizador_cds')
def otimizador_cds():
    return render_template('otimizador_cds.html')

@app.route('/p_medianas')
def p_medianas():
    return render_template('p_medianas.html')

@app.route('/pmediana_simples')
def pmediana_simples():
    return render_template('pmediana_simples.html')

# --- ROTA PARA GERAR O TEMPLATE NO FORMATO ORIGINAL ---
@app.route('/gerar_template', methods=['POST'])
def gerar_template():
    try:
        # Pega os dados do formulário
        num_cds = int(request.form.get('num_cds', 0))
        num_clientes = int(request.form.get('num_clientes', 0))
        
        if num_cds <= 0 or num_clientes <= 0:
            flash('Número de CDs e Clientes deve ser maior que zero.', 'error')
            return redirect(url_for('otimizador_cds'))

        # Criar as listas de nomes (Linhas e Colunas)
        nomes_cds = [f'CD {i}' for i in range(1, num_cds + 1)]
        nomes_clientes = [f'Cliente {i}' for i in range(1, num_clientes + 1)]
        
        # Cabeçalho completo
        colunas = ['Origem / Destino'] + nomes_clientes + ['Custo Fixo', 'Capacidade']
        
        # Montar os dados das linhas
        dados = []
        for cd in nomes_cds:
            # Para cada CD: [Nome do CD, (espaços vazios pros clientes), espaço fixo, espaço cap]
            linha = [cd] + ['' for _ in nomes_clientes] + ['', '']
            dados.append(linha)
            
        # Adicionar a linha final de Demanda
        linha_demanda = ['Demanda Total'] + ['' for _ in nomes_clientes] + ['-', '-']
        dados.append(linha_demanda)
        
        # Criar o DataFrame (Tabela)
        df = pd.DataFrame(dados, columns=colunas)
        
        # Salvar na memória como Excel com Design Profissional
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            # 1. Aba Localização (dados principais)
            df.to_excel(writer, sheet_name='Localização', index=False)
            worksheet = writer.sheets['Localização']
            
            # 2. Criar aba de coordenadas dos CDs
            coords_data = [['CD', 'Latitude', 'Longitude']]  # Header
            
            for i, cd in enumerate(nomes_cds):
                # Coordenadas em branco para usuário preencher
                coords_data.append([cd, '', ''])
            
            coords_df = pd.DataFrame(coords_data[1:], columns=coords_data[0])
            coords_df.to_excel(writer, sheet_name='Coordenadas_CDs', index=False)
            
            # Aplicar formatação profissional na aba de coordenadas
            coords_worksheet = writer.sheets['Coordenadas_CDs']
            
            # --- DEFININDO ESTILOS PARA COORDENADAS ---
            coords_header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            coords_header_font = Font(color="FFFFFF", bold=True, size=12)
            coords_border = Border(
                left=Side(style='thin'), right=Side(style='thin'), 
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            coords_alignment = Alignment(horizontal='center', vertical='center')
            
            # Formatar header da aba de coordenadas
            for col_num, column_title in enumerate(['CD', 'Latitude', 'Longitude'], 1):
                cell = coords_worksheet.cell(row=1, column=col_num)
                cell.fill = coords_header_fill
                cell.font = coords_header_font
                cell.border = coords_border
                cell.alignment = coords_alignment
            
            # Ajustar largura das colunas
            coords_worksheet.column_dimensions['A'].width = 25  # CD
            coords_worksheet.column_dimensions['B'].width = 15  # Latitude  
            coords_worksheet.column_dimensions['C'].width = 15  # Longitude
            
            # Adicionar bordas nas células de dados
            for row_num in range(2, num_cds + 2):
                for col_num in range(1, 4):
                    cell = coords_worksheet.cell(row=row_num, column=col_num)
                    cell.border = coords_border
                    cell.alignment = coords_alignment
            
            # 3. Criar aba de instruções (por último)
            from openpyxl import Workbook
            instructions_worksheet = writer.book.create_sheet('Como Usar')
            
            # --- DEFININDO OS ESTILOS PARA ABA DE INSTRUÇÕES ---
            # Cor do Cabeçalho (Azul Escuro) e Fonte Branca
            instructions_header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            instructions_header_font = Font(color="FFFFFF", bold=True, size=12)
            
            # Cor de destaque (Azul Claro)
            highlight_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            instructions_bold_font = Font(bold=True, size=11)
            
            # Fonte normal
            normal_font = Font(size=10)
            
            # Alinhamento e Bordas
            center_align = Alignment(horizontal="center", vertical="center")
            left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- ADICIONAR INSTRUÇÕES NA ABA "Como Usar" ---
            instructions_data = [
                ['📋 GUIA DE USO - PLANILHA DE OTIMIZAÇÃO DE CDs', ''],
                ['', ''],
                ['🎯 OBJETIVO', 'Encontrar a configuração ótima de Centros de Distribuição que minimiza custos totais'],
                ['', ''],
                ['📝 COMO PREENCHER A PLANILHA "Localização":', ''],
                ['1. CUSTOS DE TRANSPORTE (Células Cinzas)', 'Nas células cinza, informe o custo unitário (R$/unidade) para transportar mercadorias de cada CD para cada cliente'],
                ['2. CUSTO FIXO DOS CDs (Coluna "Custo Fixo")', 'Informe o custo mensal fixo (R$) para operar cada CD candidato'],
                ['3. CAPACIDADE DOS CDs (Coluna "Capacidade")', 'Informe a capacidade máxima (unidades/mês) que cada CD pode processar'],
                ['4. DEMANDA DOS CLIENTES (Linha "Demanda Total")', 'Informe a demanda mensal (unidades) que cada cliente precisa receber'],
                ['', ''],
                ['🗺️ COORDENADAS GEOGRÁFICAS (ABA "Coordenadas_CDs"):', ''],
                ['5. LOCALIZAÇÃO DOS CDs (Opcional)', 'Na aba "Coordenadas_CDs", informe as coordenadas geográficas (latitude/longitude) de cada CD'],
                ['6. SISTEMA DE MAPAS', 'Se preencher as coordenadas e selecionar "distância" como tipo, o sistema calculará distâncias geográficas reais e gerará um mapa interativo'],
                ['7. MAPA INTERATIVO', 'O mapa mostrará CDs abertos (verde), fechados (vermelho) e áreas de atendimento com percentual de capacidade utilizada'],
                ['', ''],
                ['🎨 GUIA VISUAL DAS CORES:', ''],
                ['AZUL ESCURO (Cabeçalho)', 'Títulos das colunas - não editar'],
                ['AZUL CLARO (Primeira coluna)', 'Nomes dos CDs e linha de demanda - não editar'],
                ['CINZA CLARO (Miolo da tabela)', 'ÁREA PARA PREENCHER SEUS DADOS'],
                ['VERDE (Coordenadas)', 'Área para preencher coordenadas geográficas'],
                ['', ''],
                ['💡 EXEMPLO PRÁTICO:', ''],
                ['Para 2 CDs (SP, RJ) e 3 Clientes (A, B, C):', ''],
                ['• Custo SP → Cliente A: R$ 5,50/unidade', ''],
                ['• Custo Fixo SP: R$ 15.000/mês', ''],
                ['• Capacidade SP: 1.000 unidades/mês', ''],
                ['• Demanda Cliente A: 300 unidades/mês', ''],
                ['• Coordenada SP: (-23.5505, -46.6333) ou (-23,5505, -46,6333)', ''],
                ['', ''],
                ['⚠️ DICAS IMPORTANTES:', ''],
                ['• Use números sem formatação (ex: 15000, não R$ 15.000)', ''],
                ['• Verifique se a capacidade total ≥ demanda total', ''],
                ['• Custos de transporte devem ser por unidade de produto', ''],
                ['• O sistema decidirá quais CDs abrir e fechar automaticamente', ''],
                ['• Com coordenadas, você terá mapas interativos detalhados', ''],
                ['• Coordenadas aceitam formato: (-23.5505, -46.6333) ou (-23,5505, -46,6333)', ''],
                ['', ''],
                ['🔍 VALIDAÇÕES ESSENCIAIS:', ''],
                ['• Capacidade Total ≥ Demanda Total', 'Verifique se a soma das capacidades dos CDs atende à demanda total'],
                ['• Valores Não Negativos', 'Custos, capacidades e demandas devem ser ≥ 0'],
                ['• Coordenadas Válidas', 'Latitude: -90 a 90, Longitude: -180 a 180'],
                ['• Matriz Completa', 'Todos os CDs devem ter valores para todos os clientes'],
                ['', ''],
                ['🛠️ TROUBLESHOOTING:', ''],
                ['• Erro: "Capacidade insuficiente"', 'Aumente capacidade dos CDs ou reduza a demanda'],
                ['• Erro: "Valor inválido"', 'Verifique se há células vazias ou texto não numérico'],
                ['• Mapa não aparece', 'Verifique se preencheu a aba "Coordenadas_CDs" corretamente'],
                ['', ''],
                [' PRÓXIMOS PASSOS:', ''],
                ['1. Preencha todos os dados na aba "Localização"', ''],
                ['2. Opcional: Preencha coordenadas na aba "Coordenadas_CDs"', ''],
                ['3. Salve o arquivo Excel', ''],
                ['4. Faça upload no sistema', ''],
                ['5. Aguarde o resultado da otimização', ''],
                ['6. Se usou coordenadas, visualize o mapa interativo nos resultados', ''],
            ]
            
            # Adicionar instruções na worksheet
            for row_idx, (col1, col2) in enumerate(instructions_data, 1):
                # Coluna 1
                cell1 = instructions_worksheet.cell(row=row_idx, column=1, value=col1)
                if row_idx == 1:
                    cell1.font = instructions_header_font
                    cell1.fill = instructions_header_fill
                    cell1.alignment = center_align
                elif col1.startswith('🎯') or col1.startswith('📝') or col1.startswith('🎨') or col1.startswith('💡') or col1.startswith('⚠️') or col1.startswith('🚀'):
                    cell1.font = instructions_bold_font
                    cell1.fill = highlight_fill
                    cell1.alignment = left_align
                else:
                    cell1.font = normal_font
                    cell1.alignment = left_align
                cell1.border = thin_border
                
                # Coluna 2
                cell2 = instructions_worksheet.cell(row=row_idx, column=2, value=col2)
                cell2.font = normal_font
                cell2.alignment = left_align
                cell2.border = thin_border
            
            # Ajustar largura das colunas da aba de instruções
            instructions_worksheet.column_dimensions['A'].width = 50
            instructions_worksheet.column_dimensions['B'].width = 80
            
            # Congelar painel para facilitar navegação
            instructions_worksheet.freeze_panes = 'A2'
            
            # --- DEFININDO OS ESTILOS PARA ABA DE DADOS ---
            # Cor do Cabeçalho (Azul Escuro) e Fonte Branca
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            
            # Cor da Primeira Coluna (Azul Claro) e Negrito
            col1_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            bold_font = Font(bold=True)
            
            # Cor da Área de Digitação (Cinza bem claro para guiar o usuário)
            input_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            
            # Alinhamento Centralizado e Bordas Finas
            center_align = Alignment(horizontal="center", vertical="center")
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- APLICANDO OS ESTILOS ---
            max_row = worksheet.max_row
            max_col = worksheet.max_column

            # 1. Estilizar o Cabeçalho (Linha 1)
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = thin_border

            # 2. Estilizar o restante da tabela
            for row in range(2, max_row + 1):
                for col in range(1, max_col + 1):
                    cell = worksheet.cell(row=row, column=col)
                    cell.border = thin_border
                    cell.alignment = center_align
                    
                    # Se for a primeira coluna (Nomes dos CDs e Demanda)
                    if col == 1:
                        cell.font = bold_font
                        cell.fill = col1_fill
                    # Se for a área onde o usuário vai digitar os números
                    else:
                        cell.fill = input_fill

            # 3. Ajustar largura das colunas automaticamente
            for col in worksheet.columns:
                max_length = 0
                column = col[0].column_letter # Pega a letra da coluna (A, B, C...)
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                # Dá um respiro de tamanho (+4) para não ficar apertado
                worksheet.column_dimensions[column].width = max_length + 4
            
            # --- ADICIONAR ABA DE COORDENADAS (NOVO) ---
            coords_data = [['CD', 'Latitude', 'Longitude']]  # Header
            
            for i, cd in enumerate(nomes_cds):
                # Coordenadas exemplo (São Paulo, Rio de Janeiro, Belo Horizonte, etc.)
                coords_exemplo = [
                    (-23.5505, -46.6333),  # São Paulo
                    (-22.9068, -43.1729),  # Rio de Janeiro  
                    (-19.9167, -43.9345),  # Belo Horizonte
                    (-30.0346, -51.2177),  # Porto Alegre
                    (-8.0476, -34.8770),  # Recife
                    (-12.9714, -38.5014),  # Salvador
                    (-15.8267, -47.9218),  # Brasília
                    (-3.1190, -60.0217),  # Manaus
                    (-5.0892, -42.8016),  # Teresina
                    (-2.5307, -44.3068)    # São Luís
                ]
                coords_data.append([cd, '', ''])
            
            coords_df = pd.DataFrame(coords_data[1:], columns=coords_data[0])
            coords_df.to_excel(writer, sheet_name='Coordenadas_CDs', index=False)
            
            # Aplicar formatação profissional na aba de coordenadas
            coords_worksheet = writer.sheets['Coordenadas_CDs']
            
            # --- DEFININDO ESTILOS PARA COORDENADAS ---
            coords_header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            coords_header_font = Font(color="FFFFFF", bold=True, size=12)
            coords_border = Border(
                left=Side(style='thin'), right=Side(style='thin'), 
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            coords_alignment = Alignment(horizontal='center', vertical='center')
            
            # Formatar header da aba de coordenadas
            for col_num, column_title in enumerate(['CD', 'Latitude', 'Longitude'], 1):
                cell = coords_worksheet.cell(row=1, column=col_num)
                cell.fill = coords_header_fill
                cell.font = coords_header_font
                cell.border = coords_border
                cell.alignment = coords_alignment
            
            # Ajustar largura das colunas
            coords_worksheet.column_dimensions['A'].width = 25  # CD
            coords_worksheet.column_dimensions['B'].width = 15  # Latitude  
            coords_worksheet.column_dimensions['C'].width = 15  # Longitude
            
            # Adicionar bordas nas células de dados
            for row_num in range(2, num_cds + 2):
                for col_num in range(1, 4):
                    cell = coords_worksheet.cell(row=row_num, column=col_num)
                    cell.border = coords_border
                    cell.alignment = coords_alignment

        output.seek(0)
        
        return send_file(
            output,
            download_name="Template_Localizacao_CDs.xlsx",
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
    except Exception as e:
        flash(f'Erro ao gerar template: {str(e)}', 'error')
        return redirect(url_for('otimizador_cds'))

# --- ATUALIZAÇÃO NA ROTA DE UPLOAD ---
@app.route('/resolver', methods=['POST'])
def resolver():
    try:
        if 'arquivo_planilha' not in request.files:
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('otimizador_cds'))
        
        file = request.files['arquivo_planilha']
        
        if file.filename == '':
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('otimizador_cds'))
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Ler dados do Excel (todas as abas para coordenadas)
            df = pd.read_excel(filepath, sheet_name=None)  # Ler todas as abas
            print(f"DataFrame lido com sucesso: {len(df)} abas")
            for aba_name, aba_df in df.items():
                print(f"  Aba '{aba_name}': {aba_df.shape}")
            
            # Usar a primeira aba como principal para a matriz
            primeira_aba = list(df.keys())[0]
            df_principal = df[primeira_aba]
            
            # Obter tipo de dado do formulário
            tipo_dado = request.form.get('tipo_dado', 'custo')
            
            # Resolver o problema tradicional
            from solver_tradicional import resolver_problema_logistica
            resultado = resolver_problema_logistica(df, tipo_dado)
            
            # Limpar arquivo temporário
            os.remove(filepath)
            
            if resultado.get('status') == 'Erro':
                flash(f"Erro na otimização: {resultado.get('mensagem')}", 'error')
                return redirect(url_for('otimizador_cds'))
            
            return render_template('resultado_tradicional.html', resultado=resultado)
                
        else:
            flash('Formato de arquivo inválido. Por favor, envie um arquivo .xlsx.', 'error')
            return redirect(url_for('otimizador_cds'))
            
    except Exception as e:
        flash(f'Erro ao processar arquivo: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

def criar_excel_pcentros(df, num_cds, num_clientes, p_cds):
    """Função para criar Excel estilizado para p-Centros"""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Dados', index=False)
        
        # Aplicar estilização igual às outras ferramentas
        worksheet = writer.sheets['Dados']
        max_row = worksheet.max_row
        max_col = worksheet.max_column
        
        # Definir estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="9c27b0", end_color="9c27b0", fill_type="solid")  # Roxo para p-Centros
        center_align = Alignment(horizontal="center", vertical="center")
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), 
                           top=Side(style='thin'), bottom=Side(style='thin'))
        col1_fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
        input_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        bold_font = Font(bold=True)
        
        # 1. Estilizar cabeçalho
        for col in range(1, max_col + 1):
            cell = worksheet.cell(row=1, column=col)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border

        # 2. Estilizar o restante da tabela
        for row in range(2, max_row + 1):
            for col in range(1, max_col + 1):
                cell = worksheet.cell(row=row, column=col)
                cell.border = thin_border
                cell.alignment = center_align
                
                # Se for a primeira coluna (Nomes dos CDs e Demanda)
                if col == 1:
                    cell.font = bold_font
                    cell.fill = col1_fill
                # Se for a área onde o usuário vai digitar os números
                else:
                    cell.fill = input_fill

        # 3. Ajustar largura das colunas automaticamente
        for col in worksheet.columns:
            max_length = 0
            column = col[0].column_letter # Pega a letra da coluna (A, B, C...)
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            # Dá um respiro de tamanho (+4) para não ficar apertado
            worksheet.column_dimensions[column].width = max_length + 4
        
        # 4. Adicionar aba de coordenadas (NOVO)
        coords_data = [['CD', 'Latitude', 'Longitude']]
        for i in range(1, num_cds + 1):
            coords_data.append([f'CD {i}', '', ''])  # Coordenadas em branco para preencher
        
        coords_df = pd.DataFrame(coords_data[1:], columns=coords_data[0])
        coords_df.to_excel(writer, sheet_name='Coordenadas_CDs', index=False)
        
        # Aplicar formatação profissional na aba de coordenadas
        coords_worksheet = writer.sheets['Coordenadas_CDs']
        
        # Estilizar cabeçalho das coordenadas
        header_fill_coords = PatternFill(start_color="4CAF50", end_color="4CAF50", fill_type="solid")  # Verde
        for col in range(1, 4):  # 3 colunas: CD, Latitude, Longitude
            cell = coords_worksheet.cell(row=1, column=col)
            cell.font = header_font
            cell.fill = header_fill_coords
            cell.alignment = center_align
            cell.border = thin_border
        
        # Estilizar células de dados das coordenadas
        for row in range(2, num_cds + 2):  # +2 por causa do cabeçalho
            for col in range(1, 4):
                cell = coords_worksheet.cell(row=row, column=col)
                cell.border = thin_border
                cell.alignment = center_align
                cell.fill = input_fill
        
        # Ajustar largura das colunas das coordenadas
        coords_worksheet.column_dimensions['A'].width = 15  # CD
        coords_worksheet.column_dimensions['B'].width = 20  # Latitude  
        coords_worksheet.column_dimensions['C'].width = 20  # Longitude
        
        # 5. Adicionar aba de instruções
        instrucoes = pd.DataFrame([
            ['Como Usar', ''],
            ['', ''],
            ['1. Preencha os dados:', ''],
            ['- Aba "Dados": Cabeçalho e matriz de distâncias', ''],
            ['- Aba "Coordenadas_CDs": Latitude e Longitude dos CDs (OPCIONAL)', ''],
            ['- Linha 1: Cabeçalho com IDs (Cliente 1, Cliente 2, ...)', ''],
            ['- Linhas seguintes: CDs (CD 1, CD 2, ...)', ''],
            ['- Valores: Distâncias entre CDs e clientes', ''],
            ['- Diagonal CD×CD: Mantenha zero (0)', ''],
            ['', ''],
            ['1.1. Coordenadas Geográficas (Opcional):', ''],
            ['- Se preenchidas, o sistema usará distâncias reais', ''],
            ['- Formato: Latitude (ex: -23.5505), Longitude (ex: -46.6333)', ''],
            ['- Use coordenadas Google Maps para precisão máxima', ''],
            ['', ''],
            ['2. Salve o arquivo:', ''],
            ['- Clique em Arquivo > Salvar Como', ''],
            ['- Escolha formato .xlsx', ''],
            ['', ''],
            ['3. Faça upload:', ''],
            ['- Volte para a página do sistema', ''],
            ['- Clique em "Escolher arquivo" e selecione a planilha', ''],
            ['- Clique em "Otimizar" para calcular os p-Centros', ''],
            ['', ''],
            ['Observações:', ''],
            ['- p-Centros minimiza a MAIOR distância', ''],
            ['- Garante equidade no atendimento', ''],
            ['- Ideal para serviços emergenciais', ''],
            ['- Sem Demanda e sem Custo Fixo', ''],
            ['', ''],
            ['🌍 Recursos Geográficos:', ''],
            ['- Com coordenadas: Mapa interativo e distâncias reais', ''],
            ['- Sem coordenadas: Usa matriz de distâncias original', ''],
            ['- Mapa mostra CDs selecionados e áreas de cobertura', ''],
            ['', ''],
            ['Parâmetros configurados:', ''],
            [f'- CDs candidatos: {num_cds}', ''],
            [f'- Clientes: {num_clientes}', ''],
            [f'- CDs a abrir: {p_cds}']
        ])
        instrucoes.to_excel(writer, sheet_name='Como Usar', index=False)
    
    return output

@app.route('/exportar_resultados_pcentros')
def exportar_resultados_pcentros():
    try:
        # Obter resultado da sessão ou usar dados mock para teste
        resultado = {
            "tipo_valor": "distancia",
            "valor_maximo": 45.5,
            "num_cds_selecionados": 2,
            "p": 2,
            "cds_selecionados": ["CD 1", "CD 3"],
            "atribuicoes": [
                {"cd": "CD 1", "cliente": "Cliente 1", "valor": 25.3},
                {"cd": "CD 1", "cliente": "Cliente 2", "valor": 18.7},
                {"cd": "CD 3", "cliente": "Cliente 3", "valor": 45.5},
                {"cd": "CD 3", "cliente": "Cliente 4", "valor": 32.1}
            ],
            "clientes_por_cd": {"CD 1": 2, "CD 3": 2}
        }
        
        # Criar conteúdo CSV
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Cabeçalho
        writer.writerow(['CD Selecionado', 'Cliente Atendido', f'{resultado["tipo_valor"].title()} (d_ij)', 'Status'])
        
        # Dados
        for atribuicao in resultado['atribuicoes']:
            status = "Máximo" if atribuicao['valor'] == resultado['valor_maximo'] else "Normal"
            writer.writerow([
                atribuicao['cd'],
                atribuicao['cliente'],
                f"{atribuicao['valor']:.2f}",
                status
            ])
        
        # Criar resposta
        output.seek(0)
        response = Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={
                'Content-Disposition': f'attachment; filename=resultado_pcentros_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
            }
        )
        
        return response
        
    except Exception as e:
        flash(f'Erro ao exportar resultados: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

@app.route('/exportar_resultados_pmedianas')
def exportar_resultados_pmedianas():
    try:
        # Obter resultado da sessão ou usar dados mock para teste
        resultado = {
            "tipo_valor": "distancia",
            "custo_total_ponderado": 1250.75,
            "num_cds_selecionados": 3,
            "p": 3,
            "total_clientes": 10,
            "cds_selecionados": ["CD 1", "CD 2", "CD 3"],
            "clientes_por_cd": {"CD 1": 4, "CD 2": 3, "CD 3": 3},
            "atribuicoes": [
                {"cd": "CD 1", "cliente": "Cliente 1", "distancia": 10.5, "demanda_atendida": 100, "custo_ponderado": 1050.0},
                {"cd": "CD 2", "cliente": "Cliente 2", "distancia": 15.5, "demanda_atendida": 80, "custo_ponderado": 1240.0},
                {"cd": "CD 3", "cliente": "Cliente 3", "distancia": 8.2, "demanda_atendida": 120, "custo_ponderado": 984.0}
            ]
        }
        
        # Criar conteúdo CSV
        csv_content = "Relatório de Resultados - p-Medianas\n\n"
        csv_content += f"Tipo de Análise,{resultado.get('tipo_valor', 'distância')}\n"
        csv_content += f"Custo Total Ponderado,{resultado.get('custo_total_ponderado', 0)}\n"
        csv_content += f"CDs Selecionados,{resultado.get('num_cds_selecionados', 0)}/{resultado.get('p', 0)}\n"
        csv_content += f"Total Clientes,{resultado.get('total_clientes', 0)}\n\n"
        
        csv_content += "CDs Selecionados\n"
        csv_content += "CD,Clientes,Status\n"
        for cd in resultado.get('cds_selecionados', []):
            clientes_cd = resultado.get('clientes_por_cd', {}).get(cd, 0)
            csv_content += f"{cd},{clientes_cd},Ativo\n"
        
        csv_content += "\nAtribuições\n"
        csv_content += "CD,Cliente,Distância,Quantidade,Custo Ponderado,Status\n"
        for atrib in resultado.get('atribuicoes', []):
            csv_content += f"{atrib.get('cd', '')},{atrib.get('cliente', '')},{atrib.get('distancia', 0)},{atrib.get('demanda_atendida', 0)},{atrib.get('custo_ponderado', 0)},Atendido\n"
        
        # Criar arquivo para download
        output = io.BytesIO()
        output.write(csv_content.encode('utf-8'))
        output.seek(0)
        
        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name='resultado_pmedianas.csv'
        )
        
    except Exception as e:
        flash(f'Erro ao exportar resultados: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

# --- ROTA PARA MÁXIMA COBERTURA ---
@app.route('/max_cobertura')
def max_cobertura():
    return render_template('max_cobertura.html')

@app.route('/gerar_template_maxcobertura', methods=['POST'])
def gerar_template_maxcobertura():
    try:
        num_cds = int(request.form.get('num_cds', 0))
        num_clientes = int(request.form.get('num_clientes', 0))
        p_cds = int(request.form.get('p_cds', 0))
        raio_cobertura = float(request.form.get('raio_cobertura', 0))
        tipo_dado = request.form.get('tipo_dado', 'distancia')
        
        if num_cds <= 0 or num_clientes <= 0 or p_cds <= 0:
            flash('Todos os valores devem ser maiores que zero.', 'error')
            return redirect(url_for('max_cobertura'))
        
        if p_cds > num_cds:
            flash('O número de CDs a abrir (p) não pode ser maior que o número de CDs candidatos.', 'error')
            return redirect(url_for('max_cobertura'))
        
        # Criar as listas de nomes
        nomes_cds = [f'CD {i}' for i in range(1, num_cds + 1)]
        nomes_clientes = [f'Cliente {i}' for i in range(1, num_clientes + 1)]
        
        # Cabeçalho completo
        colunas = ['Origem / Destino'] + nomes_clientes
        
        # Montar os dados das linhas
        dados = []
        for cd in nomes_cds:
            # Para cada CD: [Nome do CD, (espaços vazios pros clientes)]
            linha = [cd] + ['' for _ in nomes_clientes]
            dados.append(linha)
        
        # Adicionar a linha final de Demanda
        linha_demanda = ['Demanda Total'] + ['' for _ in nomes_clientes]
        dados.append(linha_demanda)
        
        # Criar o DataFrame
        df = pd.DataFrame(dados, columns=colunas)
        
        # Salvar na memória como Excel com Design Profissional
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Maxima_Cobertura', index=False)
            worksheet = writer.sheets['Maxima_Cobertura']
            
            # Criar aba de coordenadas dos CDs (NOVO)
            coords_data = [['CD', 'Latitude', 'Longitude']]  # Header
            
            for i, cd in enumerate(nomes_cds):
                # Coordenadas em branco para usuário preencher
                coords_data.append([cd, '', ''])
            
            coords_df = pd.DataFrame(coords_data[1:], columns=coords_data[0])
            coords_df.to_excel(writer, sheet_name='Coordenadas_CDs', index=False)
            
            # Aplicar formatação profissional na aba de coordenadas
            coords_worksheet = writer.sheets['Coordenadas_CDs']
            
            # --- DEFININDO ESTILOS PARA COORDENADAS ---
            coords_header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            coords_header_font = Font(color="FFFFFF", bold=True, size=12)
            coords_border = Border(
                left=Side(style='thin'), right=Side(style='thin'), 
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            coords_alignment = Alignment(horizontal='center', vertical='center')
            
            # Formatar header da aba de coordenadas
            for col_num, column_title in enumerate(['CD', 'Latitude', 'Longitude'], 1):
                cell = coords_worksheet.cell(row=1, column=col_num)
                cell.fill = coords_header_fill
                cell.font = coords_header_font
                cell.border = coords_border
                cell.alignment = coords_alignment
            
            # Ajustar largura das colunas
            coords_worksheet.column_dimensions['A'].width = 25  # CD
            coords_worksheet.column_dimensions['B'].width = 15  # Latitude  
            coords_worksheet.column_dimensions['C'].width = 15  # Longitude
            
            # Adicionar bordas nas células de dados
            for row_num in range(2, len(nomes_cds) + 2):
                for col_num in range(1, 4):
                    cell = coords_worksheet.cell(row=row_num, column=col_num)
                    cell.border = coords_border
                    cell.alignment = coords_alignment
            
            # Criar aba de instruções
            instructions_worksheet = writer.book.create_sheet('Como Usar')
            
            # --- DEFININDO OS ESTILOS PARA ABA DE INSTRUÇÕES ---
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True, size=12)
            highlight_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            bold_font = Font(bold=True, size=11)
            normal_font = Font(size=10)
            center_align = Alignment(horizontal="center", vertical="center")
            left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- ADICIONAR INSTRUÇÕES NA ABA "Como Usar" ---
            instructions_data = [
                ['📋 GUIA DE USO - MÁXIMA COBERTURA', ''],
                ['', ''],
                ['🎯 OBJETIVO', f'Abrir exatamente {p_cds} CDs para maximizar a demanda coberta dentro do {"raio de " + str(raio_cobertura) + " km" if raio_cobertura > 0 else "matriz binária"}'],
                ['', ''],
                ['📝 COMO PREENCHER A PLANILHA "Maxima_Cobertura":', ''],
                ['1. DISTÂNCIAS/CUSTOS (Células Cinzas)', f'Informe {"distâncias (km, tempo)" if tipo_dado == "distancia" else "custos (R$/unidade)"} de cada CD para cada cliente'],
                ['2. DEMANDA DOS CLIENTES (Linha "Demanda Total")', 'Informe a demanda mensal (unidades) de cada cliente'],
                ['3. MODO BINÁRIO (Opcional)', 'Se não tiver valores exatos, use 1 (coberto) ou 0 (não coberto). Neste caso, defina raio=0 no sistema.'],
                ['', ''],
                ['🎨 GUIA VISUAL DAS CORES:', ''],
                ['AZUL ESCURO (Cabeçalho)', 'Títulos das colunas - não editar'],
                ['AZUL CLARO (Primeira coluna)', 'Nomes dos CDs e linha de demanda - não editar'],
                ['CINZA CLARO (Miolo da tabela)', 'ÁREA PARA PREENCHER SEUS DADOS'],
                ['VERDE CLARO (Linha de demanda)', 'ÁREA PARA PREENCHER A DEMANDA'],
                ['', ''],
                ['💡 CARACTERÍSTICAS DO MODELO:', ''],
                ['• Foco em Cobertura', 'Maximiza o número de clientes atendidos'],
                ['• Raio de Cobertura', f'Apenas clientes dentro de {raio_cobertura} km são considerados cobertos'],
                ['• Demanda Ponderada', 'Clientes com maior demanda têm mais peso'],
                ['• Número Fixo de CDs', f'Exatamente {p_cds} CDs serão abertos'],
                ['• Atendimento Exclusivo', 'Cada cliente é atendido por apenas 1 CD'],
                ['', ''],
                ['⚠️ DICAS IMPORTANTES:', ''],
                ['• Use números sem formatação', f'Ex: 15.5 para {"distância" if tipo_dado == "distancia" else "custo"}, não {"15,5 km" if tipo_dado == "distancia" else "R$ 15,50"}'],
                ['• Raio de Cobertura', f'Configure {raio_cobertura} km no sistema para filtrar clientes'],
                ['• Demanda é Importante', 'Clientes com maior demanda priorizam a otimização'],
                ['• Coordenadas aceitam formato: (-23.5505, -46.6333) ou (-23,5505, -46,6333)', ''],
                ['', ''],
                ['🔍 VALIDAÇÕES ESSENCIAIS:', ''],
                ['• Raio de Cobertura', f'Verifique se {raio_cobertura} km é adequado para seu negócio'],
                ['• Demanda Não Negativa', 'Todos os valores de demanda devem ser ≥ 0'],
                ['• Matriz Completa', 'Todos os CDs devem ter valores para todos os clientes'],
                ['• Coordenadas Válidas', 'Latitude: -90 a 90, Longitude: -180 a 180'],
                ['', ''],
                ['🛠️ TROUBLESHOOTING:', ''],
                ['• Erro: "Nenhum cliente coberto"', 'Aumente o raio de cobertura ou verifique as distâncias'],
                ['• Erro: "Valor inválido"', 'Verifique se há células vazias ou texto não numérico'],
                ['• Baixa cobertura', 'Reduza o raio ou aumente o número de CDs a abrir'],
                ['', ''],
                ['🌍 COORDENADAS DOS CDs (OPCIONAL):', ''],
                ['• Aba "Coordenadas_CDs"', 'Coordenadas geográficas reais dos CDs com design profissional'],
                ['• Layout Profissional', 'Cabeçalho azul escuro, bordas e alinhamento centralizado'],
                ['• Latitude/Longitude', 'Use coordenadas decimais (ex: -23.5505, -46.6333)'],
                ['• Formatação Automática', 'Sistema aplica cores e estilos automaticamente'],
                ['• Precisão Melhorada', 'Sistema usará distâncias reais se preenchido'],
                ['• Totalmente Opcional', 'Funciona normalmente sem preencher esta aba'],
                ['', ''],
                ['🎨 BENEFÍCIOS VISUAIS:', ''],
                ['• Mapa Interativo', 'Visualização geográfica da solução ótima'],
                ['• Indicadores Visuais', 'CDs selecionados vs não selecionados'],
                ['• Áreas de Cobertura', 'Círculos mostrando raio de atuação'],
                ['• Legenda Integrada', 'Identificação clara dos elementos do mapa'],
                ['', ''],
                ['🚀 PRÓXIMOS PASSOS:', ''],
                ['1. Preencha todos os dados na aba "Maxima_Cobertura"', ''],
                ['2. (Opcional) Atualize coordenadas na aba "Coordenadas_CDs"', ''],
                ['3. Salve o arquivo Excel', ''],
                ['4. Faça upload no sistema', ''],
                ['5. Aguarde o resultado da otimização', ''],
            ]
            
            # Adicionar instruções na worksheet
            for row_idx, (col1, col2) in enumerate(instructions_data, 1):
                cell1 = instructions_worksheet.cell(row=row_idx, column=1, value=col1)
                if row_idx == 1:
                    cell1.font = header_font
                    cell1.fill = header_fill
                    cell1.alignment = center_align
                elif col1.startswith('🎯') or col1.startswith('📝') or col1.startswith('🎨') or col1.startswith('💡') or col1.startswith('⚠️') or col1.startswith('🚀'):
                    cell1.font = bold_font
                    cell1.fill = highlight_fill
                    cell1.alignment = left_align
                else:
                    cell1.font = normal_font
                    cell1.alignment = left_align
                cell1.border = thin_border
                
                cell2 = instructions_worksheet.cell(row=row_idx, column=2, value=col2)
                cell2.font = normal_font
                cell2.alignment = left_align
                cell2.border = thin_border
            
            # Ajustar largura das colunas da aba de instruções
            instructions_worksheet.column_dimensions['A'].width = 50
            instructions_worksheet.column_dimensions['B'].width = 80
            instructions_worksheet.freeze_panes = 'A2'
            
            # --- DEFININDO OS ESTILOS PARA ABA DE DADOS ---
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            col1_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            bold_font = Font(bold=True)
            input_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            demanda_fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
            center_align = Alignment(horizontal="center", vertical="center")
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- APLICANDO OS ESTILOS NA ABA DE DADOS ---
            max_row = worksheet.max_row
            max_col = worksheet.max_column

            # 1. Estilizar o Cabeçalho (Linha 1)
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = thin_border
            
            # 2. Estilizar a primeira coluna (nomes dos CDs e linha de demanda)
            for row in range(2, max_row):
                cell = worksheet.cell(row=row, column=1)
                cell.fill = col1_fill
                cell.font = bold_font
                cell.alignment = center_align
                cell.border = thin_border
            
            # 3. Estilizar a linha de demanda (última linha)
            demanda_row = max_row
            for col in range(1, max_col + 1):
                cell = worksheet.cell(row=demanda_row, column=col)
                cell.fill = demanda_fill
                cell.font = bold_font
                cell.alignment = center_align
                cell.border = thin_border
            
            # 4. Estilizar as células de entrada (dados)
            for row in range(2, max_row):
                for col in range(2, max_col + 1):
                    cell = worksheet.cell(row=row, column=col)
                    cell.fill = input_fill
                    cell.border = thin_border
            
            # 5. Ajustar largura das colunas
            for col in range(1, max_col + 1):
                col_letter = get_column_letter(col)
                if col == 1:
                    worksheet.column_dimensions[col_letter].width = 20
                else:
                    worksheet.column_dimensions[col_letter].width = 15
        
        output.seek(0)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='Template_MaximaCobertura.xlsx'
        )
    except Exception as e:
        flash(f'Erro ao gerar template: {str(e)}', 'error')
        return redirect(url_for('max_cobertura'))

@app.route('/resolver_maxcobertura', methods=['POST'])
def resolver_maxcobertura_route():
    try:
        print(f"=== DEBUG - ROTA MÁXIMA COBERTURA ===")
        
        # Puxando as variáveis corretas do HTML (mesmo padrão do p-Centros)
        p = int(request.form.get('p_cds', 0))
        raio = float(request.form.get('raio_cobertura', 0))
        tipo_dado = request.form.get('tipo_dado', 'distancia')
        
        print(f"Parâmetros recebidos: p={p}, raio={raio}, tipo={tipo_dado}")
        
        file = request.files.get('file')
        print(f"Arquivo recebido: {file}")
        print(f"Nome do arquivo: {file.filename if file else 'None'}")
        
        if p <= 0:
            flash('Número de CDs a abrir deve ser maior que zero.', 'error')
            return redirect(url_for('max_cobertura'))
            
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            print(f"Arquivo salvo em: {filepath}")
            
            df = pd.read_excel(filepath, sheet_name=None)  # Ler todas as abas
            print(f"DataFrame lido com sucesso: {len(df)} abas")
            for aba_name, aba_df in df.items():
                print(f"  Aba '{aba_name}': {aba_df.shape}")
            
            from solver_maxcobertura import resolver_maxcobertura
            resultado = resolver_maxcobertura(df, p, raio, tipo_dado)
            
            print(f"Resultado do solver: {resultado}")
            
            os.remove(filepath)
            
            if resultado.get('status') == 'Erro':
                flash(f"Erro na otimização: {resultado.get('mensagem')}", 'error')
                return redirect(url_for('max_cobertura'))
            
            # Adicionar mensagem de sucesso
            if resultado.get('status') == 'Sucesso':
                cds_selecionados = resultado.get('cds_selecionados', [])
                demanda_coberta = resultado.get('demanda_coberta', 0)
                percentual = resultado.get('percentual_cobertura', 0)
                mensagem = f"✅ Otimização concluída com sucesso! {len(cds_selecionados)} CDs selecionados cobrindo {percentual:.1f}% da demanda total ({demanda_coberta:.0f} unidades)."
                flash(mensagem, 'success')
                print(f"✅ Mensagem de sucesso configurada: {mensagem}")
                print(f"✅ Redirecionando para template resultado_maxcobertura.html")
                
            return render_template('resultado_maxcobertura.html', resultado=resultado)
        else:
            print("❌ Arquivo inválido ou não recebido")
            flash('Arquivo inválido ou não selecionado.', 'error')
            return redirect(url_for('max_cobertura'))
    except Exception as e:
        print(f"❌ Erro na rota: {str(e)}")
        flash(f'Erro ao processar: {str(e)}', 'error')
        return redirect(url_for('max_cobertura'))

# --- ROTA PARA p-CENTROS ---
@app.route('/p_centros')
def p_centros():
    return render_template('p_centros.html')

@app.route('/gerar_template_pcentros', methods=['POST'])
def gerar_template_pcentros():
    try:
        num_cds = int(request.form.get('num_cds', 0))
        num_clientes = int(request.form.get('num_clientes', 0))
        p_cds = int(request.form.get('p_cds', 0))
        tipo_dado = request.form.get('tipo_dado', 'distancia')
        
        if num_cds <= 0 or num_clientes <= 0 or p_cds <= 0:
            flash('Todos os valores devem ser maiores que zero.', 'error')
            return redirect(url_for('p_medianas'))
        
        if p_cds > num_cds:
            flash('O número de CDs a abrir (p) não pode ser maior que o número de CDs candidatos.', 'error')
            return redirect(url_for('p_medianas'))
        
        # Criar as listas de nomes
        nomes_cds = [f'CD {i}' for i in range(1, num_cds + 1)]
        nomes_clientes = [f'Cliente {i}' for i in range(1, num_clientes + 1)]
        
        # Cabeçalho completo (apenas distâncias/custos)
        colunas = ['Origem / Destino'] + nomes_clientes
        
        # Montar os dados das linhas
        dados = []
        for cd in nomes_cds:
            # Para cada CD: [Nome do CD, (espaços vazios pros clientes)]
            linha = [cd] + ['' for _ in nomes_clientes]
            dados.append(linha)
        
        # Criar o DataFrame
        df = pd.DataFrame(dados, columns=colunas)
        
        # Salvar na memória como Excel com Design Profissional
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Localização', index=False)
            worksheet = writer.sheets['Localização']
            
            # Criar aba de coordenadas dos CDs (NOVO)
            coords_data = [['CD', 'Latitude', 'Longitude']]  # Header
            
            for i, cd in enumerate(nomes_cds):
                # Coordenadas em branco para usuário preencher
                coords_data.append([cd, '', ''])
            
            coords_df = pd.DataFrame(coords_data[1:], columns=coords_data[0])
            coords_df.to_excel(writer, sheet_name='Coordenadas_CDs', index=False)
            
            # Aplicar formatação profissional na aba de coordenadas
            coords_worksheet = writer.sheets['Coordenadas_CDs']
            
            # --- DEFININDO ESTILOS PARA COORDENADAS ---
            coords_header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            coords_header_font = Font(color="FFFFFF", bold=True, size=12)
            coords_border = Border(
                left=Side(style='thin'), right=Side(style='thin'), 
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            coords_alignment = Alignment(horizontal='center', vertical='center')
            
            # Formatar header da aba de coordenadas
            for col_num, column_title in enumerate(['CD', 'Latitude', 'Longitude'], 1):
                cell = coords_worksheet.cell(row=1, column=col_num)
                cell.fill = coords_header_fill
                cell.font = coords_header_font
                cell.border = coords_border
                cell.alignment = coords_alignment
            
            # Ajustar largura das colunas
            coords_worksheet.column_dimensions['A'].width = 25  # CD
            coords_worksheet.column_dimensions['B'].width = 15  # Latitude  
            coords_worksheet.column_dimensions['C'].width = 15  # Longitude
            
            # Adicionar bordas nas células de dados
            for row_num in range(2, num_cds + 2):
                for col_num in range(1, 4):
                    cell = coords_worksheet.cell(row=row_num, column=col_num)
                    cell.border = coords_border
                    cell.alignment = coords_alignment
            
            # Criar aba de instruções
            instructions_worksheet = writer.book.create_sheet('Como Usar')
            
            # --- DEFININDO OS ESTILOS PARA ABA DE INSTRUÇÕES ---
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True, size=12)
            highlight_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            bold_font = Font(bold=True, size=11)
            normal_font = Font(size=10)
            center_align = Alignment(horizontal="center", vertical="center")
            left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- ADICIONAR INSTRUÇÕES NA ABA "Como Usar" ---
            instructions_data = [
                ['📋 GUIA DE USO - P-CENTROS', ''],
                ['', ''],
                ['🎯 OBJETIVO', f'Abrir exatamente {p_cds} CDs para minimizar a maior {"distância" if tipo_dado == "distancia" else "custo"} entre qualquer cliente e seu CD'],
                ['', ''],
                ['📝 COMO PREENCHER A PLANILHA "Dados":', ''],
                ['1. DISTÂNCIAS/CUSTOS (Células Cinzas)', f'Informe {"distâncias (km, tempo)" if tipo_dado == "distancia" else "custos (R$/unidade)"} de cada CD para cada cliente'],
                ['2. MATRIZ SIMPLIFICADA', 'No p-Centros não usamos Demanda, Custo Fixo ou Capacidade. Preencha apenas a matriz de valores.'],
                ['', ''],
                ['🌍 COORDENADAS DOS CDs (OPCIONAL):', ''],
                ['• Aba "Coordenadas_CDs"', 'Coordenadas geográficas reais dos CDs com design profissional'],
                ['• Layout Profissional', 'Cabeçalho azul escuro, bordas e alinhamento centralizado'],
                ['• Latitude/Longitude', 'Use coordenadas decimais (ex: -23.5505, -46.6333)'],
                ['• Formatação Automática', 'Sistema aplica cores e estilos automaticamente'],
                ['• Precisão Melhorada', 'Sistema usará distâncias reais se preenchido'],
                ['• Totalmente Opcional', 'Funciona normalmente sem preencher esta aba'],
                ['', ''],
                ['🎨 GUIA VISUAL DAS CORES:', ''],
                ['AZUL ESCURO (Cabeçalho)', 'Títulos das colunas - não editar'],
                ['AZUL CLARO (Primeira coluna)', 'Nomes dos CDs - não editar'],
                ['CINZA CLARO (Miolo da tabela)', 'ÁREA PARA PREENCHER SEUS DADOS'],
                ['', ''],
                ['💡 CARACTERÍSTICAS DO MODELO:', ''],
                ['• Foco em Equidade', 'Minimiza o pior caso (cliente mais distante)'],
                ['• Sem Demanda', 'Todos os clientes têm igual importância'],
                ['• Sem Custo Fixo', 'Não há custo para abrir CDs'],
                ['• Sem Capacidade', 'CDs podem atender clientes ilimitados'],
                ['• Atendimento Exclusivo', 'Cada cliente é atendido por apenas 1 CD'],
                ['', ''],
                ['⚠️ DICAS IMPORTANTES:', ''],
                ['• Use números sem formatação', f'Ex: 15.5 para {"distância" if tipo_dado == "distancia" else "custo"}, não {"15,5 km" if tipo_dado == "distancia" else "R$ 15,50"}'],
                ['• Foco no Pior Cenário', 'O sistema garante que nenhum cliente fique muito longe'],
                ['• Ideal para Serviços Críticos', 'Perfeito para emergências, segurança, etc.'],
                ['• Coordenadas aceitam formato: (-23.5505, -46.6333) ou (-23,5505, -46,6333)', ''],
                ['', ''],
                ['🔍 VALIDAÇÕES ESSENCIAIS:', ''],
                ['• Matriz Completa', 'Todos os CDs devem ter valores para todos os clientes'],
                ['• Valores Não Negativos', 'Distâncias/custos devem ser ≥ 0'],
                ['• Coordenadas Válidas', 'Latitude: -90 a 90, Longitude: -180 a 180'],
                ['• Número Fixo de CDs', f'Exatamente {p_cds} CDs serão abertos'],
                ['', ''],
                ['🛠️ TROUBLESHOOTING:', ''],
                ['• Erro: "Solução não encontrada"', 'Verifique se as distâncias/custos são realistas'],
                ['• Erro: "Valor inválido"', 'Verifique se há células vazias ou texto não numérico'],
                ['• Mapa não aparece', 'Verifique se preencheu a aba "Coordenadas_CDs" corretamente'],
                ['• Resultado inesperado', 'Verifique se a matriz está completa e sem erros'],
                ['', ''],
                ['🎨 BENEFÍCIOS VISUAIS:', ''],
                ['• Mapa Interativo', 'Visualização geográfica da solução ótima'],
                ['• Indicadores Visuais', 'CDs selecionados vs não selecionados'],
                ['• Áreas de Cobertura', 'Círculos mostrando raio de atuação'],
                ['• Legenda Integrada', 'Identificação clara dos elementos do mapa'],
                ['', ''],
                ['🚀 PRÓXIMOS PASSOS:', ''],
                ['1. Preencha todos os dados na aba "Dados"', ''],
                ['2. (Opcional) Preencha coordenadas na aba "Coordenadas_CDs"', ''],
                ['3. Salve o arquivo Excel', ''],
                ['4. Faça upload no sistema', ''],
                ['5. Aguarde o resultado da otimização', ''],
                ['', ''],
                ['Parâmetros configurados:', ''],
                [f'- CDs candidatos: {num_cds}', ''],
                [f'- Clientes: {num_clientes}', ''],
                [f'- CDs a abrir: {p_cds}', '']
            ]
            
            # Adicionar instruções na worksheet
            for row_idx, (col1, col2) in enumerate(instructions_data, 1):
                cell1 = instructions_worksheet.cell(row=row_idx, column=1, value=col1)
                if row_idx == 1:
                    cell1.font = header_font
                    cell1.fill = header_fill
                    cell1.alignment = center_align
                elif col1.startswith('🎯') or col1.startswith('📝') or col1.startswith('🎨') or col1.startswith('💡') or col1.startswith('⚠️') or col1.startswith('🚀'):
                    cell1.font = bold_font
                    cell1.fill = highlight_fill
                    cell1.alignment = left_align
                else:
                    cell1.font = normal_font
                    cell1.alignment = left_align
                cell1.border = thin_border
                
                cell2 = instructions_worksheet.cell(row=row_idx, column=2, value=col2)
                cell2.font = normal_font
                cell2.alignment = left_align
                cell2.border = thin_border
            
            # Ajustar largura das colunas da aba de instruções
            instructions_worksheet.column_dimensions['A'].width = 50
            instructions_worksheet.column_dimensions['B'].width = 80
            instructions_worksheet.freeze_panes = 'A2'
            
            # --- DEFININDO OS ESTILOS PARA ABA DE DADOS ---
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            col1_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            bold_font = Font(bold=True)
            input_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            center_align = Alignment(horizontal="center", vertical="center")
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- APLICANDO OS ESTILOS NA ABA DE DADOS ---
            max_row = worksheet.max_row
            max_col = worksheet.max_column

            # 1. Estilizar o Cabeçalho (Linha 1)
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = thin_border
            
            # 2. Estilizar a primeira coluna (nomes dos CDs)
            for row in range(2, max_row + 1):
                cell = worksheet.cell(row=row, column=1)
                cell.fill = col1_fill
                cell.font = bold_font
                cell.alignment = center_align
                cell.border = thin_border
            
            # 3. Estilizar as células de entrada (dados)
            for row in range(2, max_row + 1):
                for col in range(2, max_col + 1):
                    cell = worksheet.cell(row=row, column=col)
                    cell.fill = input_fill
                    cell.border = thin_border
            
            # 4. Ajustar largura das colunas
            for col in range(1, max_col + 1):
                col_letter = get_column_letter(col)
                if col == 1:
                    worksheet.column_dimensions[col_letter].width = 20
                else:
                    worksheet.column_dimensions[col_letter].width = 15
        
        output.seek(0)
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name='Template_P-Centros.xlsx'
        )
        
    except Exception as e:
        flash(f'Erro ao gerar template: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

@app.route('/resolver_pcentros', methods=['POST'])
def resolver_pcentros_route():
    try:
        if 'file' not in request.files:
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('p_medianas'))
        
        file = request.files['file']
        
        if file.filename == '':
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('p_medianas'))
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Ler dados do Excel (todas as abas para coordenadas)
            df = pd.read_excel(filepath, sheet_name=None)  # Ler todas as abas
            print(f"DataFrame lido com sucesso: {len(df)} abas")
            for aba_name, aba_df in df.items():
                print(f"  Aba '{aba_name}': {aba_df.shape}")
            
            # Obter valores do formulário
            p_cds = int(request.form.get('p_cds', 1))
            tipo_dado = request.form.get('tipo_dado', 'distancia')
            
            # Resolver o problema p-Centros
            from solver_pcentros import resolver_pcentros
            resultado = resolver_pcentros(df, p_cds, tipo_dado)
            
            # Limpar arquivo temporário
            os.remove(filepath)
            
            if resultado.get('status') == 'Erro':
                flash(f"Erro na otimização: {resultado.get('mensagem')}", 'error')
                return redirect(url_for('p_medianas'))
            
            return render_template('resultado_pcentros.html', resultado=resultado)
                
        else:
            flash('Formato de arquivo inválido. Por favor, envie um arquivo .xlsx ou .csv.', 'error')
            return redirect(url_for('p_medianas'))
            
    except Exception as e:
        flash(f'Erro ao processar arquivo: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

@app.route('/salvar_resultados_pcentros', methods=['POST'])
def salvar_resultados_pcentros():
    """Salva os resultados de uma análise p-Centros no banco de dados"""
    try:
        # Obter dados do formulário
        nome_projeto = request.form.get('nome_projeto', '').strip()
        
        if not nome_projeto:
            return jsonify({'success': False, 'message': 'Por favor, informe um nome para o projeto.'})
        
        # Obter resultados da sessão (precisamos passar os resultados via formulário)
        resultados_json = request.form.get('resultados')
        if not resultados_json:
            return jsonify({'success': False, 'message': 'Resultados não encontrados.'})
        
        import json
        try:
            print(f"DEBUG: JSON recebido p-Centros: {repr(resultados_json)}")
            resultado = json.loads(resultados_json)
        except json.JSONDecodeError as e:
            print(f"Erro no JSON p-Centros: {str(e)}")
            print(f"JSON completo p-Centros: {repr(resultados_json)}")
            return jsonify({'success': False, 'message': f'Erro nos dados do resultado: Formato JSON inválido. Detalhes: {str(e)}'})
        
        # Preparar parâmetros para salvar
        parametros = {
            'p': resultado.get('p', 0),
            'tipo_valor': resultado.get('tipo_valor', 'distancia'),
            'num_cds_selecionados': resultado.get('num_cds_selecionados', 0),
            'coordenadas_usadas': resultado.get('coordenadas_usadas', False)
        }
        
        # Salvar no banco de dados
        projeto_id = salvar_projeto(
            nome=nome_projeto,
            tipo_analise='p_centros',
            parametros=parametros,
            resultados=resultado,
            mapa_html=resultado.get('mapa_html')
        )
        
        if projeto_id:
            return jsonify({
                'success': True, 
                'message': 'Projeto salvo com sucesso!',
                'projeto_id': projeto_id
            })
        else:
            return jsonify({'success': False, 'message': 'Erro ao salvar o projeto.'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/salvar_resultados_pmedianas', methods=['POST'])
def salvar_resultados_pmedianas():
    """Salva os resultados de uma análise p-Medianas no banco de dados"""
    try:
        # Obter dados do formulário
        nome_projeto = request.form.get('nome_projeto', '').strip()
        
        if not nome_projeto:
            return jsonify({'success': False, 'message': 'Por favor, informe um nome para o projeto.'})
        
        # Obter resultados da sessão (precisamos passar os resultados via formulário)
        resultados_json = request.form.get('resultados')
        if not resultados_json:
            return jsonify({'success': False, 'message': 'Resultados não encontrados.'})
        
        import json
        resultado = json.loads(resultados_json)
        
        # Preparar parâmetros para salvar
        parametros = {
            'p': resultado.get('p', 0),
            'tipo_valor': resultado.get('tipo_valor', 'distancia'),
            'num_cds_selecionados': resultado.get('num_cds_selecionados', 0),
            'custo_total': resultado.get('custo_total', 0),
            'coordenadas_usadas': resultado.get('coordenadas_usadas', False)
        }
        
        # Salvar no banco de dados
        projeto_id = salvar_projeto(
            nome=nome_projeto,
            tipo_analise='p_medianas',
            parametros=parametros,
            resultados=resultado,
            mapa_html=resultado.get('mapa_html')
        )
        
        if projeto_id:
            return jsonify({
                'success': True, 
                'message': 'Projeto salvo com sucesso!',
                'projeto_id': projeto_id
            })
        else:
            return jsonify({'success': False, 'message': 'Erro ao salvar o projeto.'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/salvar_resultados_tradicional', methods=['POST'])
def salvar_resultados_tradicional():
    """Salva os resultados de uma análise Tradicional no banco de dados"""
    try:
        # Obter dados do formulário
        nome_projeto = request.form.get('nome_projeto', '').strip()
        
        if not nome_projeto:
            return jsonify({'success': False, 'message': 'Por favor, informe um nome para o projeto.'})
        
        # Obter resultados da sessão (precisamos passar os resultados via formulário)
        resultados_json = request.form.get('resultados')
        if not resultados_json:
            return jsonify({'success': False, 'message': 'Resultados não encontrados.'})
        
        # Debug: mostrar o JSON recebido
        print(f"JSON recebido Tradicional (primeiros 200 chars): {resultados_json[:200]}")
        
        import json
        try:
            resultado = json.loads(resultados_json)
        except json.JSONDecodeError as e:
            print(f"Erro no JSON Tradicional: {str(e)}")
            print(f"JSON completo Tradicional: {repr(resultados_json)}")
            return jsonify({'success': False, 'message': f'Erro nos dados: Formato JSON inválido. Detalhes: {str(e)}'})
        
        # Preparar parâmetros para salvar
        parametros = {
            'tipo_valor': resultado.get('tipo_valor', 'distancia'),
            'num_cds_abertos': resultado.get('cds_abertos', 0) if resultado.get('cds_abertos') else 0,
            'custo_total': resultado.get('custo_total', 0),
            'coordenadas_usadas': resultado.get('coordenadas_usadas', False)
        }
        
        # Salvar no banco de dados
        projeto_id = salvar_projeto(
            nome=nome_projeto,
            tipo_analise='tradicional',
            parametros=parametros,
            resultados=resultado,
            mapa_html=resultado.get('mapa_html')
        )
        
        if projeto_id:
            return jsonify({
                'success': True, 
                'message': 'Projeto salvo com sucesso!',
                'projeto_id': projeto_id
            })
        else:
            return jsonify({'success': False, 'message': 'Erro ao salvar o projeto.'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/salvar_resultados_maxcobertura', methods=['POST'])
def salvar_resultados_maxcobertura():
    """Salva os resultados de uma análise Max Cobertura no banco de dados"""
    try:
        # Obter dados do formulário
        nome_projeto = request.form.get('nome_projeto', '').strip()
        
        if not nome_projeto:
            return jsonify({'success': False, 'message': 'Por favor, informe um nome para o projeto.'})
        
        # Obter resultados da sessão (precisamos passar os resultados via formulário)
        resultados_json = request.form.get('resultados')
        if not resultados_json:
            return jsonify({'success': False, 'message': 'Resultados não encontrados.'})
        
        # Debug: mostrar o JSON recebido
        print(f"JSON recebido Max Cobertura (primeiros 200 chars): {resultados_json[:200]}")
        
        import json
        try:
            resultado = json.loads(resultados_json)
        except json.JSONDecodeError as e:
            print(f"Erro no JSON Max Cobertura: {str(e)}")
            print(f"JSON completo Max Cobertura: {repr(resultados_json)}")
            return jsonify({'success': False, 'message': f'Erro nos dados: Formato JSON inválido. Detalhes: {str(e)}'})
        
        # Preparar parâmetros para salvar
        parametros = {
            'p': resultado.get('p', 0),
            'tipo_valor': resultado.get('tipo_valor', 'distancia'),
            'num_cds_selecionados': resultado.get('num_cds_selecionados', 0),
            'cobertura_total': resultado.get('cobertura_total', 0),
            'coordenadas_usadas': resultado.get('coordenadas_usadas', False)
        }
        
        # Salvar no banco de dados
        projeto_id = salvar_projeto(
            nome=nome_projeto,
            tipo_analise='max_cobertura',
            parametros=parametros,
            resultados=resultado,
            mapa_html=resultado.get('mapa_html')
        )
        
        if projeto_id:
            return jsonify({
                'success': True, 
                'message': 'Projeto salvo com sucesso!',
                'projeto_id': projeto_id
            })
        else:
            return jsonify({'success': False, 'message': 'Erro ao salvar o projeto.'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro: {str(e)}'})

@app.route('/ver_projeto/<int:projeto_id>')
def ver_projeto(projeto_id):
    """Visualiza um projeto salvo"""
    projeto = carregar_resultados_projeto(projeto_id)
    if not projeto:
        flash('Projeto não encontrado.', 'error')
        return redirect(url_for('dashboard'))
    
    # Renderizar com base no tipo de análise
    if projeto['tipo_analise'] == 'p_centros':
        return render_template('resultado_pcentros.html', resultado=projeto['resultados'])
    elif projeto['tipo_analise'] == 'p_medianas':
        return render_template('resultado_pmedianas.html', resultado=projeto['resultados'])
    elif projeto['tipo_analise'] == 'max_cobertura':
        return render_template('resultado_maxcobertura.html', resultado=projeto['resultados'])
    else:
        return render_template('resultado_tradicional.html', resultado=projeto['resultados'])

@app.route('/excluir_projeto/<int:projeto_id>', methods=['POST'])
def excluir_projeto_route(projeto_id):
    """Exclui um projeto"""
    success = excluir_projeto(projeto_id)
    if success:
        return jsonify({'success': True, 'message': 'Projeto excluído com sucesso!'})
    else:
        return jsonify({'success': False, 'message': 'Erro ao excluir o projeto.'})

@app.route('/relatorio_projeto/<int:projeto_id>')
def relatorio_projeto(projeto_id):
    """Gera um relatório em PDF/Excel de um projeto"""
    projeto = carregar_resultados_projeto(projeto_id)
    if not projeto:
        flash('Projeto não encontrado.', 'error')
        return redirect(url_for('dashboard'))
    
    # Por enquanto, redireciona para a exportação padrão
    # Futuro: implementar geração de PDF/Excel personalizado
    if projeto['tipo_analise'] == 'p_centros':
        return redirect(url_for('exportar_resultados_pcentros'))
    elif projeto['tipo_analise'] == 'p_medianas':
        return redirect(url_for('exportar_resultados_pmedianas'))
    elif projeto['tipo_analise'] == 'max_cobertura':
        return redirect(url_for('exportar_resultados_maxcobertura'))
    else:
        return redirect(url_for('exportar_resultados_tradicional'))

@app.route('/gerar_relatorio_anual', methods=['POST'])
def gerar_relatorio_anual():
    """Gera relatório anual dos projetos selecionados"""
    try:
        data = request.get_json()
        project_ids = data.get('project_ids', [])
        
        if not project_ids:
            return jsonify({'success': False, 'message': 'Nenhum projeto selecionado.'})
        
        # Carregar projetos selecionados
        projetos_selecionados = []
        for project_id in project_ids:
            projeto = carregar_resultados_projeto(int(project_id))
            if projeto:
                projetos_selecionados.append(projeto)
        
        if not projetos_selecionados:
            return jsonify({'success': False, 'message': 'Nenhum projeto válido encontrado.'})
        
        # Preparar dados para o template
        from datetime import datetime
        ano_atual = datetime.now().year
        data_atual = datetime.now().strftime('%d/%m/%Y')
        
        # Contar tipos de análises
        tipos_contagem = {}
        for projeto in projetos_selecionados:
            tipo = projeto['tipo_analise']
            tipos_contagem[tipo] = tipos_contagem.get(tipo, 0) + 1
        
        # Determinar período de análise
        if projetos_selecionados:
            datas = [p.get('data_criacao') for p in projetos_selecionados if p.get('data_criacao')]
            if datas:
                datas_ordenadas = sorted(datas)
                periodo_analise = f"{datas_ordenadas[0][5:7]}/{datas_ordenadas[0][:4]} - {datas_ordenadas[-1][5:7]}/{datas_ordenadas[-1][:4]}"
            else:
                periodo_analise = "N/A"
        else:
            periodo_analise = "N/A"
        
        # Funções auxiliares para o template
        def get_tipo_nome(tipo):
            nomes = {
                'p_centros': 'p-Centros',
                'p_medianas': 'p-Medianas', 
                'max_cobertura': 'Máxima Cobertura',
                'tradicional': 'Tradicional'
            }
            return nomes.get(tipo, 'Desconhecido')
        
        def get_tipo_icone(tipo):
            icones = {
                'p_centros': 'center_focus_strong',
                'p_medianas': 'location_city',
                'max_cobertura': 'radar',
                'tradicional': 'inventory_2'
            }
            return icones.get(tipo, 'help')
        
        def get_tipo_descricao(tipo):
            descricoes = {
                'p_centros': 'Minimização da distância máxima de atendimento',
                'p_medianas': 'Otimização com número fixo de instalações',
                'max_cobertura': 'Maximização da demanda dentro do raio de cobertura',
                'tradicional': 'Localização com custos fixos e capacidades'
            }
            return descricoes.get(tipo, 'Análise logística')
        
        def formatar_data(data_str):
            if not data_str:
                return 'N/A'
            return f"{data_str[8:10]}/{data_str[5:7]}/{data_str[:4]}"
        
        def formatar_data_completa(data_str):
            if not data_str:
                return 'N/A'
            return f"{data_str[8:10]}/{data_str[5:7]}/{data_str[:4]} {data_str[11:16]}"
        
        def format_currency_brl(value):
            if not value:
                return "0,00"
            return f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        
        # Função para calcular métricas consolidadas
        def calcular_consolidados(projetos):
            total_cds = 0
            total_clientes = 0
            total_investimento = 0.0
            
            for projeto in projetos:
                resultados = projeto.get('resultados', {})
                tipo = projeto.get('tipo_analise')
                
                if tipo == 'tradicional':
                    total_cds += len(resultados.get('cds_abertos', []))
                    # Contar clientes únicos em vez de rotas
                    clientes_unicos = set()
                    transportes = resultados.get('transportes', [])
                    for transporte in transportes:
                        if transporte.get('destino'):
                            clientes_unicos.add(transporte['destino'])
                    total_clientes += len(clientes_unicos)
                    total_investimento += float(resultados.get('custo_total', 0))
                    
                elif tipo == 'p_medianas':
                    total_cds += resultados.get('num_cds_selecionados', 0)
                    total_clientes += len(resultados.get('atribuicoes', []))
                    total_investimento += float(resultados.get('custo_total_ponderado', 0))
                    
                elif tipo == 'p_centros':
                    total_cds += resultados.get('num_cds_selecionados', 0)
                    total_clientes += len(resultados.get('atribuicoes', []))
                    # p-centros não tem custo monetário
                    
                elif tipo == 'max_cobertura':
                    # Máxima Cobertura: múltiplas formas de obter número de CDs
                    cds_selecionados = 0
                    if resultados.get('num_cds_selecionados'):
                        cds_selecionados = resultados.get('num_cds_selecionados')
                    elif resultados.get('cds_selecionados'):
                        cds_selecionados = len(resultados.get('cds_selecionados'))
                    elif resultados.get('atribuicoes'):
                        # Extrair CDs únicos das atribuições
                        cds_unicos = set()
                        for atribuicao in resultados.get('atribuicoes', []):
                            if atribuicao.get('cd_selecionado'):
                                cds_unicos.add(atribuicao.get('cd_selecionado'))
                        cds_selecionados = len(cds_unicos)
                    
                    total_cds += cds_selecionados
                    
                    # Contar clientes cobertos
                    clientes_cobertos = len([a for a in resultados.get('atribuicoes', []) if a.get('coberto')])
                    total_clientes += clientes_cobertos
                    # Máxima cobertura não tem custo monetário
            
            return {
                'total_cds': total_cds,
                'total_clientes': total_clientes,
                'total_investimento': total_investimento
            }
        
        # Calcular métricas consolidadas
        consolidados = calcular_consolidados(projetos_selecionados)
        
        # Função para processar dados dos projetos e garantir campos necessários
        def processar_dados_projetos(projetos):
            for projeto in projetos:
                resultados = projeto.get('resultados', {})
                tipo = projeto.get('tipo_analise')
                
                if tipo == 'tradicional' and resultados:
                    # Garantir volume_por_cd para cálculo de porcentagens
                    if 'volume_por_cd' not in resultados and resultados.get('transportes'):
                        volume_por_cd = {}
                        for transporte in resultados.get('transportes', []):
                            cd_origem = transporte.get('origem')
                            quantidade = transporte.get('quantidade', 0)
                            if cd_origem in volume_por_cd:
                                volume_por_cd[cd_origem] += quantidade
                            else:
                                volume_por_cd[cd_origem] = quantidade
                        resultados['volume_por_cd'] = volume_por_cd
                
                elif tipo == 'max_cobertura' and resultados:
                    # Garantir campos para Máxima Cobertura
                    if 'num_cds_selecionados' not in resultados:
                        if resultados.get('cds_selecionados'):
                            resultados['num_cds_selecionados'] = len(resultados.get('cds_selecionados'))
                        elif resultados.get('atribuicoes'):
                            cds_unicos = set()
                            for atribuicao in resultados.get('atribuicoes', []):
                                if atribuicao.get('cd_selecionado'):
                                    cds_unicos.add(atribuicao.get('cd_selecionado'))
                            resultados['num_cds_selecionados'] = len(cds_unicos)
                    
                    if 'clientes_atendidos' not in resultados:
                        clientes_cobertos = len([a for a in resultados.get('atribuicoes', []) if a.get('coberto')])
                        resultados['clientes_atendidos'] = clientes_cobertos
                    
                    if 'demanda_nao_coberta' not in resultados:
                        demanda_nao_coberta = 0
                        for atribuicao in resultados.get('atribuicoes', []):
                            if not atribuicao.get('coberto') and atribuicao.get('demanda'):
                                demanda_nao_coberta += atribuicao.get('demanda', 0)
                        resultados['demanda_nao_coberta'] = demanda_nao_coberta
        
        # Processar dados dos projetos
        processar_dados_projetos(projetos_selecionados)
        
        # Renderizar template corrigido
        html_content = render_template('relatorio.html', 
                                     projetos=projetos_selecionados,
                                     total_cds_consolidados=consolidados['total_cds'],
                                     total_clientes_consolidados=consolidados['total_clientes'],
                                     total_investimento_consolidado=consolidados['total_investimento'],
                                     ano_atual=ano_atual,
                                     data_atual=data_atual,
                                     tipos_contagem=tipos_contagem,
                                     periodo_analise=periodo_analise,
                                     get_tipo_nome=get_tipo_nome,
                                     get_tipo_icone=get_tipo_icone,
                                     get_tipo_descricao=get_tipo_descricao,
                                     formatar_data=formatar_data,
                                     formatar_data_completa=formatar_data_completa,
                                     format_currency_brl=format_currency_brl)
        
        return jsonify({
            'success': True, 
            'html_content': html_content,
            'message': f'Relatório gerado com {len(projetos_selecionados)} projetos!'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erro ao gerar relatório: {str(e)}'})

@app.route('/relatorio_anual_visualizacao')
def relatorio_anual_visualizacao():
    """Visualização do relatório anual (para teste)"""
    try:
        # Carregar todos os projetos do banco de dados
        projetos = listar_projetos()
        
        if not projetos:
            flash('Nenhum projeto encontrado para gerar relatório.', 'warning')
            return redirect(url_for('dashboard'))
        
        # Preparar dados (mesma lógica da rota POST)
        from datetime import datetime
        ano_atual = datetime.now().year
        data_atual = datetime.now().strftime('%d/%m/%Y')
        
        tipos_contagem = {}
        for projeto in projetos:
            tipo = projeto['tipo_analise']
            tipos_contagem[tipo] = tipos_contagem.get(tipo, 0) + 1
        
        datas = [p.get('data_criacao') for p in projetos if p.get('data_criacao')]
        if datas:
            datas_ordenadas = sorted(datas)
            periodo_analise = f"{datas_ordenadas[0][5:7]}/{datas_ordenadas[0][:4]} - {datas_ordenadas[-1][5:7]}/{datas_ordenadas[-1][:4]}"
        else:
            periodo_analise = "N/A"
        
        def get_tipo_nome(tipo):
            nomes = {
                'p_centros': 'p-Centros',
                'p_medianas': 'p-Medianas', 
                'max_cobertura': 'Máxima Cobertura',
                'tradicional': 'Tradicional'
            }
            return nomes.get(tipo, 'Desconhecido')
        
        def get_tipo_icone(tipo):
            icones = {
                'p_centros': 'center_focus_strong',
                'p_medianas': 'location_city',
                'max_cobertura': 'radar',
                'tradicional': 'inventory_2'
            }
            return icones.get(tipo, 'help')
        
        def get_tipo_descricao(tipo):
            descricoes = {
                'p_centros': 'Minimização da distância máxima de atendimento',
                'p_medianas': 'Otimização com número fixo de instalações',
                'max_cobertura': 'Maximização da demanda dentro do raio de cobertura',
                'tradicional': 'Localização com custos fixos e capacidades'
            }
            return descricoes.get(tipo, 'Análise logística')
        
        def formatar_data(data_str):
            if not data_str:
                return 'N/A'
            return f"{data_str[8:10]}/{data_str[5:7]}/{data_str[:4]}"
        
        def formatar_data_completa(data_str):
            if not data_str:
                return 'N/A'
            return f"{data_str[8:10]}/{data_str[5:7]}/{data_str[:4]} {data_str[11:16]}"
        
        def format_currency_brl(value):
            if not value:
                return "0,00"
            return f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        
        # Função para calcular métricas consolidadas
        def calcular_consolidados(projetos):
            total_cds = 0
            total_clientes = 0
            total_investimento = 0.0
            
            for projeto in projetos:
                resultados = projeto.get('resultados', {})
                tipo = projeto.get('tipo_analise')
                
                if tipo == 'tradicional':
                    total_cds += len(resultados.get('cds_abertos', []))
                    # Contar clientes únicos em vez de rotas
                    clientes_unicos = set()
                    transportes = resultados.get('transportes', [])
                    for transporte in transportes:
                        if transporte.get('destino'):
                            clientes_unicos.add(transporte['destino'])
                    total_clientes += len(clientes_unicos)
                    total_investimento += float(resultados.get('custo_total', 0))
                    
                elif tipo == 'p_medianas':
                    total_cds += resultados.get('num_cds_selecionados', 0)
                    total_clientes += len(resultados.get('atribuicoes', []))
                    total_investimento += float(resultados.get('custo_total_ponderado', 0))
                    
                elif tipo == 'p_centros':
                    total_cds += resultados.get('num_cds_selecionados', 0)
                    total_clientes += len(resultados.get('atribuicoes', []))
                    # p-centros não tem custo monetário
                    
                elif tipo == 'max_cobertura':
                    # Máxima Cobertura: múltiplas formas de obter número de CDs
                    cds_selecionados = 0
                    if resultados.get('num_cds_selecionados'):
                        cds_selecionados = resultados.get('num_cds_selecionados')
                    elif resultados.get('cds_selecionados'):
                        cds_selecionados = len(resultados.get('cds_selecionados'))
                    elif resultados.get('atribuicoes'):
                        # Extrair CDs únicos das atribuições
                        cds_unicos = set()
                        for atribuicao in resultados.get('atribuicoes', []):
                            if atribuicao.get('cd_selecionado'):
                                cds_unicos.add(atribuicao.get('cd_selecionado'))
                        cds_selecionados = len(cds_unicos)
                    
                    total_cds += cds_selecionados
                    
                    # Contar clientes cobertos
                    clientes_cobertos = len([a for a in resultados.get('atribuicoes', []) if a.get('coberto')])
                    total_clientes += clientes_cobertos
                    # Máxima cobertura não tem custo monetário
            
            return {
                'total_cds': total_cds,
                'total_clientes': total_clientes,
                'total_investimento': total_investimento
            }
        
        # Calcular métricas consolidadas
        consolidados = calcular_consolidados(projetos)
        
        # Função para processar dados dos projetos e garantir campos necessários
        def processar_dados_projetos(projetos):
            for projeto in projetos:
                resultados = projeto.get('resultados', {})
                tipo = projeto.get('tipo_analise')
                
                if tipo == 'tradicional' and resultados:
                    # Garantir volume_por_cd para cálculo de porcentagens
                    if 'volume_por_cd' not in resultados and resultados.get('transportes'):
                        volume_por_cd = {}
                        for transporte in resultados.get('transportes', []):
                            cd_origem = transporte.get('origem')
                            quantidade = transporte.get('quantidade', 0)
                            if cd_origem in volume_por_cd:
                                volume_por_cd[cd_origem] += quantidade
                            else:
                                volume_por_cd[cd_origem] = quantidade
                        resultados['volume_por_cd'] = volume_por_cd
                
                elif tipo == 'max_cobertura' and resultados:
                    # Garantir campos para Máxima Cobertura
                    if 'num_cds_selecionados' not in resultados:
                        if resultados.get('cds_selecionados'):
                            resultados['num_cds_selecionados'] = len(resultados.get('cds_selecionados'))
                        elif resultados.get('atribuicoes'):
                            cds_unicos = set()
                            for atribuicao in resultados.get('atribuicoes', []):
                                if atribuicao.get('cd_selecionado'):
                                    cds_unicos.add(atribuicao.get('cd_selecionado'))
                            resultados['num_cds_selecionados'] = len(cds_unicos)
                    
                    if 'clientes_atendidos' not in resultados:
                        clientes_cobertos = len([a for a in resultados.get('atribuicoes', []) if a.get('coberto')])
                        resultados['clientes_atendidos'] = clientes_cobertos
                    
                    if 'demanda_nao_coberta' not in resultados:
                        demanda_nao_coberta = 0
                        for atribuicao in resultados.get('atribuicoes', []):
                            if not atribuicao.get('coberto') and atribuicao.get('demanda'):
                                demanda_nao_coberta += atribuicao.get('demanda', 0)
                        resultados['demanda_nao_coberta'] = demanda_nao_coberta
        
        # Processar dados dos projetos
        processar_dados_projetos(projetos)
        
        return render_template('relatorio.html', 
                             projetos=projetos,
                             total_cds_consolidados=consolidados['total_cds'],
                             total_clientes_consolidados=consolidados['total_clientes'],
                             total_investimento_consolidado=consolidados['total_investimento'],
                             ano_atual=ano_atual,
                             data_atual=data_atual,
                             tipos_contagem=tipos_contagem,
                             periodo_analise=periodo_analise,
                             get_tipo_nome=get_tipo_nome,
                             get_tipo_icone=get_tipo_icone,
                             get_tipo_descricao=get_tipo_descricao,
                             formatar_data=formatar_data,
                             formatar_data_completa=formatar_data_completa,
                             format_currency_brl=format_currency_brl)
        
    except Exception as e:
        flash(f'Erro ao carregar relatório: {str(e)}', 'error')
        return redirect(url_for('dashboard'))

@app.route('/exportar_resultados_maxcobertura')
def exportar_resultados_maxcobertura():
    try:
        # Obter resultado da sessão ou usar dados mock para teste
        resultado = {
            "tipo_valor": "distancia",
            "demanda_coberta": 8500.0,
            "percentual_cobertura": 85.0,
            "p": 3,
            "raio_cobertura": 50,
            "cds_selecionados": ["CD 1", "CD 2", "CD 3"],
            "atribuicoes": [
                {"cd": "CD 1", "cliente": "Cliente 1", "coberto": True, "distancia": 45.0, "demanda": 1000},
                {"cd": "CD 2", "cliente": "Cliente 2", "coberto": True, "distancia": 38.0, "demanda": 800},
                {"cd": "CD 3", "cliente": "Cliente 3", "coberto": False, "distancia": 75.0, "demanda": 500}
            ]
        }
        
        # Criar conteúdo CSV
        csv_content = "Relatório de Resultados - Máxima Cobertura\n\n"
        csv_content += f"Tipo de Análise,{resultado.get('tipo_valor', 'distância')}\n"
        csv_content += f"Demanda Total Coberta,{resultado.get('demanda_coberta', 0)}\n"
        csv_content += f"Percentual de Cobertura,{resultado.get('percentual_cobertura', 0)}%\n"
        csv_content += f"CDs Selecionados,{len(resultado.get('cds_selecionados', []))}/{resultado.get('p', 0)}\n"
        csv_content += f"Raio de Cobertura,{resultado.get('raio_cobertura', 0)} km\n\n"
        
        csv_content += "CDs Selecionados\n"
        csv_content += "CD,Status,Raio\n"
        for cd in resultado.get('cds_selecionados', []):
            csv_content += f"{cd},Ativo,{resultado.get('raio_cobertura', 0)} km\n"
        
        csv_content += "\nStatus dos Clientes\n"
        csv_content += "Cliente,Status,CD Atendente,Distância,Demanda\n"
        for atrib in resultado.get('atribuicoes', []):
            status = "Coberto" if atrib.get('coberto') else "Não Coberto"
            csv_content += f"{atrib.get('cliente', '')},{status},{atrib.get('cd', '')},{atrib.get('distancia', 0)},{atrib.get('demanda', 0)}\n"
        
        # Criar arquivo para download
        output = io.BytesIO()
        output.write(csv_content.encode('utf-8'))
        output.seek(0)
        
        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name='resultado_maxcobertura.csv'
        )
        
    except Exception as e:
        flash(f'Erro ao exportar resultados: {str(e)}', 'error')
        return redirect(url_for('max_cobertura'))

@app.route('/mapa/<filename>')
def servir_mapa(filename):
    """Serve o arquivo HTML do mapa gerado"""
    try:
        # Construir o caminho completo do arquivo na pasta estática
        mapa_path = os.path.join(app.static_folder, 'mapas', filename)
        
        if os.path.exists(mapa_path):
            return send_file(
                mapa_path,
                mimetype='text/html',
                as_attachment=False,
                download_name=filename
            )
        else:
            print(f"Arquivo do mapa não encontrado: {mapa_path}")
            flash('Arquivo do mapa não encontrado.', 'error')
            return redirect(url_for('max_cobertura'))
            
    except Exception as e:
        print(f"Erro ao carregar mapa: {str(e)}")
        flash(f'Erro ao carregar mapa: {str(e)}', 'error')
        return redirect(url_for('max_cobertura'))

@app.route('/exportar_resultados_tradicional')
def exportar_resultados_tradicional():
    try:
        # Obter resultado da sessão ou usar dados mock para teste
        resultado = {
            "tipo_valor": "custo",
            "custo_total": 15000.50,
            "cds_abertos": ["CD 1", "CD 2"],
            "rotas": [
                {"cd_origem": "CD 1", "cd_destino": "Cliente A", "quantidade": 100, "custo_unitario": 5.50, "custo_total": 550.00},
                {"cd_origem": "CD 2", "cd_destino": "Cliente B", "quantidade": 80, "custo_unitario": 7.20, "custo_total": 576.00}
            ]
        }
        
        # Criar conteúdo CSV
        csv_content = "Relatório de Resultados - Otimizador Tradicional\n\n"
        csv_content += f"Tipo de Análise,{resultado.get('tipo_valor', 'custo')}\n"
        csv_content += f"Custo Total,{resultado.get('custo_total', 0)}\n"
        csv_content += f"CDs Abertos,{len(resultado.get('cds_abertos', []))}\n\n"
        
        csv_content += "CDs Abertos\n"
        csv_content += "CD,Status\n"
        for cd in resultado.get('cds_abertos', []):
            csv_content += f"{cd},Aberto\n"
        
        csv_content += "\nRotas de Transporte\n"
        csv_content += "CD Origem,CD Destino,Quantidade,Custo Unitário,Custo Total\n"
        for rota in resultado.get('rotas', []):
            csv_content += f"{rota.get('cd_origem', '')},{rota.get('cd_destino', '')},{rota.get('quantidade', 0)},{rota.get('custo_unitario', 0)},{rota.get('custo_total', 0)}\n"
        
        # Criar arquivo para download
        output = io.BytesIO()
        output.write(csv_content.encode('utf-8'))
        output.seek(0)
        
        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name='resultado_tradicional.csv'
        )
        
    except Exception as e:
        flash(f'Erro ao exportar resultados: {str(e)}', 'error')
        return redirect(url_for('otimizador_cds'))
    try:
        # Obter resultado da sessão ou usar dados mock para teste
        resultado = {
            "tipo_valor": "distancia",
            "valor_maximo": 15.5,
            "num_cds_selecionados": 3,
            "p": 3,
            "cds_selecionados": ["CD 1", "CD 2", "CD 3"],
            "atribuicoes": [
                {"cd": "CD 1", "cliente": "Cliente 1", "valor": 10.5},
                {"cd": "CD 2", "cliente": "Cliente 2", "valor": 15.5},
                {"cd": "CD 3", "cliente": "Cliente 3", "valor": 8.2}
            ]
        }
        
        # Criar conteúdo CSV
        csv_content = "Relatório de Resultados - p-Centros\n\n"
        csv_content += f"Tipo de Análise,{resultado.get('tipo_valor', 'distância')}\n"
        csv_content += f"Valor Máximo,{resultado.get('valor_maximo', 0)}\n"
        csv_content += f"CDs Selecionados,{resultado.get('num_cds_selecionados', 0)}/{resultado.get('p', 0)}\n\n"
        
        csv_content += "CDs Selecionados\n"
        csv_content += "CD,Status\n"
        for cd in resultado.get('cds_selecionados', []):
            csv_content += f"{cd},Ativo\n"
        
        csv_content += "\nAtribuições\n"
        csv_content += "CD,Cliente,Valor,Status\n"
        for atrib in resultado.get('atribuicoes', []):
            csv_content += f"{atrib.get('cd', '')},{atrib.get('cliente', '')},{atrib.get('valor', 0)},Atendido\n"
        
        # Criar arquivo para download
        output = io.BytesIO()
        output.write(csv_content.encode('utf-8'))
        output.seek(0)
        
        return send_file(
            output,
            mimetype='text/csv',
            as_attachment=True,
            download_name='resultado_pcentros.csv'
        )
        
    except Exception as e:
        flash(f'Erro ao exportar resultados: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

# --- ROTA PARA p-MEDIANAS ---
@app.route('/gerar_template_pmedianas', methods=['POST'])
def gerar_template_pmedianas():
    try:
        # Pega os dados do formulário
        num_cds = int(request.form.get('num_cds', 0))
        num_clientes = int(request.form.get('num_clientes', 0))
        p_cds = int(request.form.get('p_cds', 0))
        
        if num_cds <= 0 or num_clientes <= 0 or p_cds <= 0:
            flash('Todos os valores devem ser maiores que zero.', 'error')
            return redirect(url_for('p_medianas'))
        
        if p_cds > num_cds:
            flash('O número de CDs a abrir (p) não pode ser maior que o número de CDs candidatos.', 'error')
            return redirect(url_for('p_medianas'))

        # Criar as listas de nomes
        nomes_cds = [f'CD {i}' for i in range(1, num_cds + 1)]
        nomes_clientes = [f'Cliente {i}' for i in range(1, num_clientes + 1)]
        
        # Cabeçalho completo (com custo fixo e capacidade para manter consistência)
        colunas = ['Origem / Destino'] + nomes_clientes + ['Custo Fixo', 'Capacidade']
        
        # Montar os dados das linhas
        dados = []
        for cd in nomes_cds:
            # Para cada CD: [Nome do CD, (espaços vazios pros clientes), espaço fixo, espaço cap]
            linha = [cd] + ['' for _ in nomes_clientes] + ['', '']
            dados.append(linha)
            
        # Adicionar a linha final de Demanda
        linha_demanda = ['Demanda Total'] + ['' for _ in nomes_clientes] + ['-', '-']
        dados.append(linha_demanda)
        
        # Criar o DataFrame
        df = pd.DataFrame(dados, columns=colunas)
        
        # Salvar na memória como Excel com Design Profissional
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Localização', index=False)
            worksheet = writer.sheets['Localização']
            
            # --- ADICIONAR ABA DE COORDENADAS (NOVO) ---
            coords_data = [['CD', 'Latitude', 'Longitude']]  # Header
            
            for i, cd in enumerate(nomes_cds):
                # Coordenadas exemplo (São Paulo, Rio de Janeiro, Belo Horizonte, etc.)
                coords_exemplo = [
                    (-23.5505, -46.6333),  # São Paulo
                    (-22.9068, -43.1729),  # Rio de Janeiro  
                    (-19.9167, -43.9345),  # Belo Horizonte
                    (-30.0346, -51.2177),  # Porto Alegre
                    (-8.0476, -34.8770),  # Recife
                    (-12.9714, -38.5014),  # Salvador
                    (-15.8267, -47.9218),  # Brasília
                    (-3.1190, -60.0217),  # Manaus
                    (-5.0892, -42.8016),  # Teresina
                    (-2.5307, -44.3068)    # São Luís
                ]
                coords_data.append([cd, '', ''])
            
            coords_df = pd.DataFrame(coords_data[1:], columns=coords_data[0])
            coords_df.to_excel(writer, sheet_name='Coordenadas_CDs', index=False)
            
            # Aplicar formatação profissional na aba de coordenadas
            coords_worksheet = writer.sheets['Coordenadas_CDs']
            
            # --- DEFININDO ESTILOS PARA COORDENADAS ---
            coords_header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            coords_header_font = Font(color="FFFFFF", bold=True, size=12)
            coords_border = Border(
                left=Side(style='thin'), right=Side(style='thin'), 
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            coords_alignment = Alignment(horizontal='center', vertical='center')
            
            # Formatar header da aba de coordenadas
            for col_num, column_title in enumerate(['CD', 'Latitude', 'Longitude'], 1):
                cell = coords_worksheet.cell(row=1, column=col_num)
                cell.fill = coords_header_fill
                cell.font = coords_header_font
                cell.border = coords_border
                cell.alignment = coords_alignment
            
            # Ajustar largura das colunas
            coords_worksheet.column_dimensions['A'].width = 25  # CD
            coords_worksheet.column_dimensions['B'].width = 15  # Latitude  
            coords_worksheet.column_dimensions['C'].width = 15  # Longitude
            
            # Adicionar bordas nas células de dados
            for row_num in range(2, num_cds + 2):
                for col_num in range(1, 4):
                    cell = coords_worksheet.cell(row=row_num, column=col_num)
                    cell.border = coords_border
                    cell.alignment = coords_alignment
            
            # Criar aba de instruções
            instructions_worksheet = writer.book.create_sheet('Como Usar')
            
            # --- DEFININDO OS ESTILOS PARA ABA DE INSTRUÇÕES ---
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True, size=12)
            highlight_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            bold_font = Font(bold=True, size=11)
            normal_font = Font(size=10)
            center_align = Alignment(horizontal="center", vertical="center")
            left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- ADICIONAR INSTRUÇÕES NA ABA "Como Usar" ---
            instructions_data = [
                ['📋 GUIA DE USO - P-MEDIANAS', ''],
                ['', ''],
                ['🎯 OBJETIVO', f'Abrir exatamente {p_cds} CDs para minimizar a distância total ponderada pela demanda'],
                ['', ''],
                ['📝 COMO PREENCHER A PLANILHA "P-Medianas":', ''],
                ['1. CUSTOS DE TRANSPORTE (Células Cinzas)', 'Nas células cinza, informe o custo unitário (R$/unidade) ou distância (km) de cada CD para cada cliente'],
                ['2. DEMANDA DOS CLIENTES (Linha "Demanda Total")', 'Informe a demanda mensal (unidades) de cada cliente'],
                ['3. CUSTO FIXO E CAPACIDADE (Colunas finais)', 'Essas colunas existem para manter formato, mas NÃO são usadas no modelo p-Medianas'],
                ['', ''],
                ['🌍 COORDENADAS DOS CDs (OPCIONAL):', ''],
                ['• Aba "Coordenadas_CDs"', 'Coordenadas geográficas reais dos CDs com design profissional'],
                ['• Layout Profissional', 'Cabeçalho azul escuro, bordas e alinhamento centralizado'],
                ['• Latitude/Longitude', 'Use coordenadas decimais (ex: -23.5505, -46.6333)'],
                ['• Formatação Automática', 'Sistema aplica cores e estilos automaticamente'],
                ['• Mapas Interativos', 'Se preenchida, gera mapas com áreas de influência'],
                ['', ''],
                ['📊 SISTEMA DE CÁLCULO:', ''],
                ['• Análise Ótima', 'Sistema encontrou solução usando matriz de distâncias original'],
                ['• Sem Coordenadas', 'Coordenadas não foram utilizadas nesta otimização'],
                ['• Resultado Valido', 'Solução ótima encontrada com sucesso'],
                ['', ''],
                ['🎨 GUIA VISUAL DAS CORES:', ''],
                ['AZUL ESCURO (Cabeçalho)', 'Títulos das colunas - não editar'],
                ['AZUL CLARO (Primeira coluna)', 'Nomes dos CDs e linha de demanda - não editar'],
                ['CINZA CLARO (Miolo da tabela)', 'ÁREA PARA PREENCHER SEUS DADOS'],
                ['', ''],
                ['💡 DIFERENÇAS PARA O MODELO ANTERIOR:', ''],
                ['• Sem Custo Fixo', 'Coluna existe mas não é usada no cálculo'],
                ['• Sem Capacidade', 'Coluna existe mas os CDs têm capacidade infinita'],
                ['• Número Fixo de CDs', f'Exatamente {p_cds} CDs serão abertos'],
                ['• Atendimento Exclusivo', 'Cada cliente é atendido por apenas 1 CD'],
                ['', ''],
                ['⚠️ DICAS IMPORTANTES:', ''],
                ['• Use números sem formatação (ex: 5.50, não R$ 5,50)', ''],
                ['• O sistema atribuirá cada cliente ao CD mais próximo', ''],
                ['• Demanda pondera a importância de cada cliente', ''],
                ['• Coordenadas aceitam formato: (-23.5505, -46.6333) ou (-23,5505, -46,6333)', ''],
                ['', ''],
                ['🔍 VALIDAÇÕES ESSENCIAIS:', ''],
                ['• Matriz Completa', 'Todos os CDs devem ter valores para todos os clientes'],
                ['• Demanda Não Negativa', 'Todos os valores de demanda devem ser ≥ 0'],
                ['• Coordenadas Válidas', 'Latitude: -90 a 90, Longitude: -180 a 180'],
                ['• Número Fixo de CDs', f'Exatamente {p_cds} CDs serão abertos'],
                ['', ''],
                ['🛠️ TROUBLESHOOTING:', ''],
                ['• Erro: "Solução não encontrada"', 'Verifique se as distâncias/custos são realistas'],
                ['• Erro: "Valor inválido"', 'Verifique se há células vazias ou texto não numérico'],
                ['• Mapa não aparece', 'Verifique se preencheu a aba "Coordenadas_CDs" corretamente'],
                ['• Resultado inesperado', 'Verifique se a matriz de distâncias está correta e completa'],
                ['', ''],
                ['🚀 PRÓXIMOS PASSOS:', ''],
                ['1. Preencha todos os dados na aba "P-Medianas"', ''],
                ['2. (Opcional) Preencha coordenadas na aba "Coordenadas_CDs"', ''],
                ['3. Salve o arquivo Excel', ''],
                ['4. Faça upload no sistema', ''],
                ['5. Aguarde o resultado da otimização', ''],
            ]
            
            # Adicionar instruções na worksheet
            for row_idx, (col1, col2) in enumerate(instructions_data, 1):
                cell1 = instructions_worksheet.cell(row=row_idx, column=1, value=col1)
                if row_idx == 1:
                    cell1.font = header_font
                    cell1.fill = header_fill
                    cell1.alignment = center_align
                elif col1.startswith('🎯') or col1.startswith('📝') or col1.startswith('🎨') or col1.startswith('💡') or col1.startswith('⚠️') or col1.startswith('🚀'):
                    cell1.font = bold_font
                    cell1.fill = highlight_fill
                    cell1.alignment = left_align
                else:
                    cell1.font = normal_font
                    cell1.alignment = left_align
                cell1.border = thin_border
                
                cell2 = instructions_worksheet.cell(row=row_idx, column=2, value=col2)
                cell2.font = normal_font
                cell2.alignment = left_align
                cell2.border = thin_border
            
            # Ajustar largura das colunas da aba de instruções
            instructions_worksheet.column_dimensions['A'].width = 50
            instructions_worksheet.column_dimensions['B'].width = 80
            instructions_worksheet.freeze_panes = 'A2'
            
            # --- DEFININDO OS ESTILOS PARA ABA DE DADOS ---
            header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            header_font = Font(color="FFFFFF", bold=True)
            col1_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            bold_font = Font(bold=True)
            input_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
            center_align = Alignment(horizontal="center", vertical="center")
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            
            # --- APLICANDO OS ESTILOS NA ABA DE DADOS ---
            max_row = worksheet.max_row
            max_col = worksheet.max_column

            # 1. Estilizar o Cabeçalho (Linha 1)
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align
                cell.border = thin_border

            # 2. Estilizar o restante da tabela
            for row in range(2, max_row + 1):
                for col in range(1, max_col + 1):
                    cell = worksheet.cell(row=row, column=col)
                    cell.border = thin_border
                    cell.alignment = center_align
                    
                    # Se for a primeira coluna (Nomes dos CDs)
                    if col == 1:
                        cell.font = bold_font
                        cell.fill = col1_fill
                    # Se for a área onde o usuário vai digitar os números
                    else:
                        cell.fill = input_fill

            # 3. Ajustar largura das colunas automaticamente
            for col in worksheet.columns:
                max_length = 0
                column = col[0].column_letter # Pega a letra da coluna (A, B, C...)
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                # Dá um respiro de tamanho (+4) para não ficar apertado
                worksheet.column_dimensions[column].width = max_length + 4
            
            # --- ADICIONAR ABA DE COORDENADAS (NOVO) ---
            coords_data = [['CD', 'Latitude', 'Longitude']]  # Header
            
            for i, cd in enumerate(nomes_cds):
                # Coordenadas exemplo (São Paulo, Rio de Janeiro, Belo Horizonte, etc.)
                coords_exemplo = [
                    (-23.5505, -46.6333),  # São Paulo
                    (-22.9068, -43.1729),  # Rio de Janeiro  
                    (-19.9167, -43.9345),  # Belo Horizonte
                    (-30.0346, -51.2177),  # Porto Alegre
                    (-8.0476, -34.8770),  # Recife
                    (-12.9714, -38.5014),  # Salvador
                    (-15.8267, -47.9218),  # Brasília
                    (-3.1190, -60.0217),  # Manaus
                    (-5.0892, -42.8016),  # Teresina
                    (-2.5307, -44.3068)    # São Luís
                ]
                coords_data.append([cd, '', ''])
            
            coords_df = pd.DataFrame(coords_data[1:], columns=coords_data[0])
            coords_df.to_excel(writer, sheet_name='Coordenadas_CDs', index=False)
            
            # Aplicar formatação profissional na aba de coordenadas
            coords_worksheet = writer.sheets['Coordenadas_CDs']
            
            # --- DEFININDO ESTILOS PARA COORDENADAS ---
            coords_header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
            coords_header_font = Font(color="FFFFFF", bold=True, size=12)
            coords_border = Border(
                left=Side(style='thin'), right=Side(style='thin'), 
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            coords_alignment = Alignment(horizontal='center', vertical='center')
            
            # Formatar header da aba de coordenadas
            for col_num, column_title in enumerate(['CD', 'Latitude', 'Longitude'], 1):
                cell = coords_worksheet.cell(row=1, column=col_num)
                cell.fill = coords_header_fill
                cell.font = coords_header_font
                cell.border = coords_border
                cell.alignment = coords_alignment
            
            # Ajustar largura das colunas
            coords_worksheet.column_dimensions['A'].width = 25  # CD
            coords_worksheet.column_dimensions['B'].width = 15  # Latitude  
            coords_worksheet.column_dimensions['C'].width = 15  # Longitude
            
            # Adicionar bordas nas células de dados
            for row_num in range(2, num_cds + 2):
                for col_num in range(1, 4):
                    cell = coords_worksheet.cell(row=row_num, column=col_num)
                    cell.border = coords_border
                    cell.alignment = coords_alignment

        output.seek(0)
        
        return send_file(
            output,
            download_name="Template_P-Medianas.xlsx",
            as_attachment=True,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
    except Exception as e:
        flash(f'Erro ao gerar template: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

@app.route('/resolver_pmedianas', methods=['POST'])
def resolver_pmedianas():
    try:
        if 'arquivo_planilha' not in request.files:
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('p_medianas'))
        
        file = request.files['arquivo_planilha']
        
        if file.filename == '':
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('p_medianas'))
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Ler dados do Excel (todas as abas para coordenadas)
            df = pd.read_excel(filepath, sheet_name=None)  # Ler todas as abas
            print(f"DataFrame lido com sucesso: {len(df)} abas")
            for aba_name, aba_df in df.items():
                print(f"  Aba '{aba_name}': {aba_df.shape}")
            
            # Extrair valores do formulário
            p_cds = int(request.form.get('p_cds', 0))
            tipo_dado = request.form.get('tipo_dado', 'distancia')
            
            # Resolver o problema p-medianas
            from solver_pmedianas import resolver_pmedianas
            resultado = resolver_pmedianas(df, p_cds, tipo_dado)
            
            # Limpar arquivo temporário
            os.remove(filepath)
            
            if resultado.get('status') == 'Erro':
                flash(f"Erro na otimização: {resultado.get('mensagem')}", 'error')
                return redirect(url_for('p_medianas'))
            
            return render_template('resultado_pmedianas.html', resultado=resultado)
                
        else:
            flash('Formato de arquivo inválido. Por favor, envie um arquivo .xlsx ou .csv.', 'error')
            return redirect(url_for('p_medianas'))
            
    except Exception as e:
        flash(f'Erro ao processar arquivo: {str(e)}', 'error')
        return redirect(url_for('p_medianas'))

@app.route('/gerar_template_pmediana_simples', methods=['POST'])
def gerar_template_pmediana_simples():
    try:
        # Pega os dados do formulário
        num_cds = int(request.form.get('num_cds', 0))
        num_clientes = int(request.form.get('num_clientes', 0))
        p_cds = int(request.form.get('p_cds', 1))
        
        if num_cds <= 0 or num_clientes <= 0:
            flash('Número de CDs e clientes deve ser maior que zero.', 'error')
            return redirect(url_for('pmediana_simples'))
        
        # Criar DataFrame para a planilha (apenas coordenadas)
        dados = []
        for i in range(num_cds):
            dados.append({
                'ID': f'P{i+1}',
                'Coordenada X': 0.0,
                'Coordenada Y': 0.0
            })
        
        df = pd.DataFrame(dados)
        
        # Criar arquivo Excel em memória
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Dados', index=False)
            
            # Adicionar aba de instruções
            instrucoes = pd.DataFrame([
                ['Como Usar', ''],
                ['', ''],
                ['1. Preencha os dados:', ''],
                ['- ID: Identificador único para cada ponto', ''],
                ['- Coordenada X: Posição X no plano cartesiano', ''],
                ['- Coordenada Y: Posição Y no plano cartesiano', ''],
                ['', ''],
                ['2. Salve o arquivo:', ''],
                ['- Clique em Arquivo > Salvar Como', ''],
                ['- Escolha formato .xlsx', ''],
                ['', ''],
                ['3. Faça upload:', ''],
                ['- Volte para a página do sistema', ''],
                ['- Clique em "Escolher arquivo" e selecione a planilha', ''],
                ['- Clique em "Calcular p-Mediana"', ''],
                ['', ''],
                ['Observações:', ''],
                ['- Mínimo de 1 ponto necessário', ''],
                ['- As coordenadas podem ser decimais', ''],
                ['- O sistema encontrará o ponto ótimo automaticamente', ''],
                ['', ''],
                ['Parâmetros configurados:', ''],
                [f'- Pontos de análise: {num_cds}', ''],
                [f'- Clientes (referência): {num_clientes}', ''],
                [f'- p-Medianas: {p_cds} (será usado para cálculo)', '']
            ])
            instrucoes.to_excel(writer, sheet_name='Como Usar', index=False)
        
        output.seek(0)
        
        return send_file(
            output,
            as_attachment=True,
            download_name=f'template_pmediana_simples_{num_cds}_pontos.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    except Exception as e:
        flash(f'Erro ao gerar template: {str(e)}', 'error')
        return redirect(url_for('pmediana_simples'))

# --- ROTA PARA p-MEDIANA SIMPLES ---
@app.route('/resolver_pmediana_simples', methods=['POST'])
def resolver_pmediana_simples_route():
    try:
        if 'arquivo_planilha' not in request.files:
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('pmediana_simples'))
        
        file = request.files['arquivo_planilha']
        
        if file.filename == '':
            flash('Nenhum arquivo foi selecionado.', 'error')
            return redirect(url_for('pmediana_simples'))
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Ler dados do Excel
            df_principal = pd.read_excel(filepath)
            
            # Resolver o problema p-Mediana simples
            from solver_pmediana_simples import resolver_pmediana_simples
            resultado = resolver_pmediana_simples(df_principal)
            
            # Limpar arquivo temporário
            os.remove(filepath)
            
            if resultado.get('status') == 'Erro':
                flash(f"Erro na otimização: {resultado.get('mensagem')}", 'error')
                return redirect(url_for('pmediana_simples'))
            
            return render_template('resultado_pmediana_simples.html', resultado=resultado)
                
        else:
            flash('Formato de arquivo inválido. Por favor, envie um arquivo .xlsx ou .csv.', 'error')
            return redirect(url_for('pmediana_simples'))
            
    except Exception as e:
        flash(f'Erro ao processar arquivo: {str(e)}', 'error')
        return redirect(url_for('pmediana_simples'))

# --- ROTA PARA GERAR RELATÓRIO EM WORD ---
@app.route('/gerar_relatorio_word')
def gerar_relatorio_word():
    """Gera o relatório anual em formato Word (.docx) com design profissional"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from datetime import datetime
        import tempfile
        import os
        
        # Carregar todos os projetos do banco de dados
        projetos = listar_projetos()
        
        if not projetos:
            flash('Nenhum projeto encontrado para gerar relatório.', 'warning')
            return redirect(url_for('dashboard'))
        
        # ESSENCIAL: Processar dados dos projetos para calcular porcentagens e métricas
        def processar_dados_projetos(projetos):
            for projeto in projetos:
                resultados = projeto.get('resultados', {})
                tipo = projeto.get('tipo_analise')
                
                if tipo == 'tradicional' and resultados:
                    # Garantir volume_por_cd para cálculo de porcentagens
                    if 'volume_por_cd' not in resultados and resultados.get('transportes'):
                        volume_por_cd = {}
                        for transporte in resultados.get('transportes', []):
                            cd_origem = transporte.get('origem')
                            quantidade = transporte.get('quantidade', 0)
                            if cd_origem in volume_por_cd:
                                volume_por_cd[cd_origem] += quantidade
                            else:
                                volume_por_cd[cd_origem] = quantidade
                        resultados['volume_por_cd'] = volume_por_cd
                    
                    # Calcular distribuição de CDs com porcentagens
                    if 'volume_por_cd' in resultados and 'distribuicao_cds' not in projeto:
                        volume_total = sum(resultados['volume_por_cd'].values())
                        distribuicao_cds = []
                        for cd, volume in resultados['volume_por_cd'].items():
                            percentual = (volume / volume_total * 100) if volume_total > 0 else 0
                            distribuicao_cds.append({
                                'nome': cd,
                                'volume': volume,
                                'percentual': percentual
                            })
                        projeto['distribuicao_cds'] = distribuicao_cds
                
                elif tipo == 'max_cobertura' and resultados:
                    # Calcular métricas de cobertura
                    atribuicoes = resultados.get('atribuicoes', [])
                    total_clientes = len(atribuicoes)
                    clientes_cobertos = len([a for a in atribuicoes if a.get('coberto')])
                    clientes_nao_cobertos = total_clientes - clientes_cobertos
                    
                    projeto['cobertura_stats'] = {
                        'total_clientes': total_clientes,
                        'clientes_cobertos': clientes_cobertos,
                        'clientes_nao_cobertos': clientes_nao_cobertos,
                        'taxa_cobertura': (clientes_cobertos / total_clientes * 100) if total_clientes > 0 else 0
                    }
                    
                    # Calcular clientes por CD para máxima cobertura
                    clientes_por_cd = {}
                    for atribuicao in atribuicoes:
                        cd = atribuicao.get('cd_selecionado', atribuicao.get('cd', 'N/A'))
                        if cd != 'N/A':
                            clientes_por_cd[cd] = clientes_por_cd.get(cd, 0) + 1
                    resultados['clientes_por_cd'] = clientes_por_cd
            
            # Garantir que p-medianas e p-centros tenham clientes_por_cd
            if tipo in ['p_medianas', 'p_centros'] and resultados:
                atribuicoes = resultados.get('atribuicoes', [])
                clientes_por_cd = {}
                for atribuicao in atribuicoes:
                    cd = atribuicao.get('cd_selecionado', atribuicao.get('cd', 'N/A'))
                    if cd != 'N/A':
                        clientes_por_cd[cd] = clientes_por_cd.get(cd, 0) + 1
                resultados['clientes_por_cd'] = clientes_por_cd
        
        # Processar os dados
        processar_dados_projetos(projetos)
        
        # Calcular métricas consolidadas
        def calcular_consolidados(projetos):
            total_cds = 0
            total_clientes = 0
            total_investimento = 0.0
            
            for projeto in projetos:
                resultados = projeto.get('resultados', {})
                tipo = projeto.get('tipo_analise')
                
                if tipo == 'tradicional':
                    total_cds += len(resultados.get('cds_abertos', []))
                    # Contar clientes únicos em vez de rotas
                    clientes_unicos = set()
                    transportes = resultados.get('transportes', [])
                    for transporte in transportes:
                        if transporte.get('destino'):
                            clientes_unicos.add(transporte['destino'])
                    total_clientes += len(clientes_unicos)
                    total_investimento += float(resultados.get('custo_total', 0))
                
                elif tipo == 'p_medianas':
                    total_cds += resultados.get('num_cds_selecionados', 0)
                    total_clientes += len(resultados.get('atribuicoes', []))
                    total_investimento += float(resultados.get('custo_total_ponderado', 0))
                
                elif tipo == 'p_centros':
                    total_cds += resultados.get('num_cds_selecionados', 0)
                    total_clientes += len(resultados.get('atribuicoes', []))
                    # p-centros não tem custo monetário
                
                elif tipo == 'max_cobertura':
                    # Máxima Cobertura: múltiplas formas de obter número de CDs
                    cds_selecionados = 0
                    if resultados.get('num_cds_selecionados'):
                        cds_selecionados = resultados.get('num_cds_selecionados')
                    elif resultados.get('cds_selecionados'):
                        cds_selecionados = len(resultados.get('cds_selecionados'))
                    elif resultados.get('atribuicoes'):
                        # Extrair CDs únicos das atribuições
                        cds_unicos = set()
                        for atribuicao in resultados.get('atribuicoes', []):
                            if atribuicao.get('cd_selecionado'):
                                cds_unicos.add(atribuicao.get('cd_selecionado'))
                        cds_selecionados = len(cds_unicos)
                    
                    total_cds += cds_selecionados
                    
                    # Contar clientes cobertos
                    clientes_cobertos = len([a for a in resultados.get('atribuicoes', []) if a.get('coberto')])
                    total_clientes += clientes_cobertos
                    # Máxima cobertura não tem custo monetário
            
            return {
                'total_cds': total_cds,
                'total_clientes': total_clientes,
                'total_investimento': total_investimento
            }
        
        # Calcular métricas consolidadas
        consolidados = calcular_consolidados(projetos)
        
        # Preparar dados
        data_atual = datetime.now().strftime('%d/%m/%Y')
        
        def get_tipo_nome(tipo):
            nomes = {
                'p_centros': 'p-Centros',
                'p_medianas': 'p-Medianas', 
                'max_cobertura': 'Máxima Cobertura',
                'tradicional': 'Tradicional'
            }
            return nomes.get(tipo, 'Desconhecido')
        
        def get_tipo_descricao(tipo):
            descricoes = {
                'p_centros': 'Minimização da distância máxima de atendimento',
                'p_medianas': 'Otimização com número fixo de instalações',
                'max_cobertura': 'Maximização da demanda dentro do raio de cobertura',
                'tradicional': 'Localização com custos fixos e capacidades'
            }
            return descricoes.get(tipo, 'Análise logística')
        
        def format_currency_brl(value):
            if not value:
                return "0,00"
            return f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        
        def get_tipo_cor(tipo):
            cores = {
                'p_centros': RGBColor(156, 39, 176),    # Roxo
                'p_medianas': RGBColor(34, 197, 94),    # Verde
                'max_cobertura': RGBColor(59, 130, 246), # Azul
                'tradicional': RGBColor(249, 115, 22)  # Laranja
            }
            return cores.get(tipo, RGBColor(37, 99, 235))  # Azul padrão
        
        # Criar documento Word
        doc = Document()
        
        # Configurar margens
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # ============= CAPA PROFISSIONAL =============
        # Adicionar cabeçalho com gradiente simulado
        title = doc.add_heading('Relatório Técnico Anual', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Formatar título
        for run in title.runs:
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
        
        # Subtítulo
        subtitle = doc.add_paragraph('Análise Otimizada de Redes Logísticas e Soluções de Distribuição')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for run in subtitle.runs:
            run.font.size = Pt(16)
            run.font.italic = True
            run.font.color.rgb = RGBColor(107, 114, 128)
        
        # Logo NexusNode
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_para.add_run('📦 NexusNode')
        logo_run.font.size = Pt(18)
        logo_run.font.bold = True
        logo_run.font.color.rgb = RGBColor(37, 99, 235)
        
        # Espaçamento
        for _ in range(3):
            doc.add_paragraph()
        
        # Tabela de informações da capa
        info_table = doc.add_table(rows=2, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Formatar células
        for row in info_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
        
        # Preencher informações
        info_table.cell(0, 0).text = 'RESPONSÁVEL'
        info_table.cell(0, 1).text = 'SISTEMA'
        info_table.cell(1, 0).text = 'EMISSÃO'
        info_table.cell(1, 1).text = data_atual
        
        # Quebra de página
        doc.add_page_break()
        
        # ============= FUNDAMENTAÇÃO TEÓRICA =============
        doc.add_heading('1. Fundamentação Teórica', level=1)
        
        # Parágrafos introdutórios
        intro1 = doc.add_paragraph(
            'O presente relatório técnico tem como objetivo apresentar uma análise abrangente dos resultados obtidos '
            'através da aplicação de modelos matemáticos de Otimização Combinatória em redes logísticas, '
            'desenvolvidos e implementados mediante a plataforma inteligente NexusNode. A crescente complexidade '
            'dos cenários modernos de distribuição exige decisões estratégicas fundamentadas em dados '
            'estruturados e análises quantitativas robustas, visando garantir máxima eficiência '
            'operacional e competitividade no mercado.'
        )
        
        intro2 = doc.add_paragraph(
            'O foco central deste documento consiste na compilação sistemática de análises detalhadas referentes '
            'à localização ótima de facilidades logísticas (Centros de Distribuição), buscando alcançar o '
            'equilíbrio ideal entre múltiplos objetivos estratégicos: redução significativa de custos '
            'operacionais, minimização de distâncias percorridas na cadeia de suprimentos e '
            'maximização da cobertura geográfica de atendimento aos pontos de demanda identificados.'
        )
        
        # Caixa de destaque
        highlight_para = doc.add_paragraph()
        highlight_run = highlight_para.add_run('Objetivo Principal: ')
        highlight_run.bold = True
        highlight_run.font.color.rgb = RGBColor(15, 23, 42)
        
        highlight_para.add_run(
            'Fornecer uma base analítica e matemática rigorosa, fundamentada em algoritmos de '
            'otimização avançados, para subsidiar o processo decisório de planejamento '
            'estratégico da cadeia de suprimentos da organização no período avaliado, '
            'possibilitando maximização de eficiência e rentabilidade.'
        )
        
        # Metodologia
        doc.add_heading('2. Metodologia de Análise', level=1)
        
        metodologia_para = doc.add_paragraph(
            'Para a obtenção dos resultados apresentados neste relatório técnico, foram empregados '
            'algoritmos computacionais avançados de Pesquisa Operacional, especificamente '
            'desenvolvidos e implementados através de abordagens matemáticas distintas e complementares, '
            'cada uma com suas características e aplicações específicas:'
        )
        
        # Descrições dos modelos utilizados
        modelos_utilizados = set(projeto['tipo_analise'] for projeto in projetos)
        
        if 'p_medianas' in modelos_utilizados:
            doc.add_heading('2.1. Modelo p-Medianas', level=2)
            doc.add_paragraph(
                'Este modelo matemático concentra-se primordialmente na otimização de eficiência de custos. '
                'O algoritmo computacional aloca estrategicamente um número predefinido (p) de '
                'instalações logísticas de forma a minimizar a distância total ponderada entre as '
                'instalações selecionadas e os pontos de demanda, considerando o custo de transporte como '
                'variável decisiva. Esta abordagem mostra-se ideal para cenários corporativos onde o '
                'custo logístico representa o fator mais crítico e impactante na rentabilidade '
                'operacional do negócio.'
            )
        
        if 'p_centros' in modelos_utilizados:
            doc.add_heading('2.2. Modelo p-Centros', level=2)
            doc.add_paragraph(
                'Com foco estratégico em equidade distributiva e atendimento urgente, o modelo '
                'p-Centros busca minimizar a distância máxima absoluta entre qualquer ponto de demanda e '
                'o Centro de Distribuição geograficamente mais próximo. Esta metodologia representa o '
                'padrão ouro para operações logísticas que requerem garantia de tempo máximo de '
                'resposta, equalizando o nível de serviço em toda a rede de distribuição e '
                'assegurando atendimento uniforme a todos os clientes.'
            )
        
        if 'max_cobertura' in modelos_utilizados:
            doc.add_heading('2.3. Modelo de Máxima Cobertura', level=2)
            doc.add_paragraph(
                'O modelo de Máxima Cobertura prioriza expansão de mercado e alcance geográfico, '
                'maximizando o número de pontos de demanda atendidos dentro de um raio de distância '
                'pré-determinado. Esta abordagem é particularmente valiosa para empresas em fase de '
                'expansão ou aquelas que buscam aumentar sua participação de mercado através da '
                'otimização estratégica da cobertura geográfica, garantindo atendimento ao maior número '
                'possível de clientes com recursos limitados.'
            )
        
        if 'tradicional' in modelos_utilizados:
            doc.add_heading('2.4. Modelo Tradicional', level=2)
            doc.add_paragraph(
                'O modelo tradicional de localização capacitada combina otimização integrada de custos '
                'fixos e variáveis, selecionando estrategicamente os Centros de Distribuição a serem '
                'abertos e alocando clientes aos CDs mais próximos considerando capacidades limitadas. '
                'Esta abordagem é fundamental para operações logísticas estabelecidas que buscam '
                'equilibrar custos operacionais fixos com custos de transporte, garantindo '
                'atendimento eficiente da demanda dentro das restrições de capacidade existentes.'
            )
        
        # ============= ANÁLISE DOS PROJETOS =============
        for i, projeto in enumerate(projetos, 1):
            doc.add_page_break()
            
            # Cabeçalho do projeto
            doc.add_heading(f'3.{i} Projeto: {projeto["nome"]}', level=1)
            
            # Parágrafo de tipo de análise
            tipo_para = doc.add_paragraph()
            tipo_para.add_run('Tipo de Análise: ').bold = True
            tipo_para.add_run(get_tipo_descricao(projeto["tipo_analise"]))
            
            # Obter tipo e resultados para uso posterior
            resultados = projeto.get('resultados', {})
            tipo = projeto.get('tipo_analise')
            
            # ============= MAPA INTERATIVO =============
            if projeto.get('mapa_html'):
                doc.add_heading('🗺️ Mapa Interativo', level=2)
                
                # Informações do mapa
                mapa_para = doc.add_paragraph()
                mapa_para.add_run('Arquivo do Mapa: ').bold = True
                mapa_para.add_run(f'{projeto["mapa_html"]}')
                
                # Adicionar detalhes sobre o mapa
                mapa_info_para = doc.add_paragraph()
                mapa_info_para.add_run('📍 Localização: ').bold = True
                mapa_info_para.add_run('Mapa geográfico interativo com visualização dos CDs e clientes')
                
                # Adicionar informações de coordenadas se disponíveis
                if projeto.get('coordenadas_info'):
                    coords_para = doc.add_paragraph()
                    coords_para.add_run('🌍 Coordenadas: ').bold = True
                    coords_para.add_run('Sistema utilizou coordenadas geográficas reais para posicionamento preciso')
                else:
                    coords_para = doc.add_paragraph()
                    coords_para.add_run('📊 Distâncias: ').bold = True
                    coords_para.add_run('Sistema utilizou matriz de distâncias para posicionamento relativo')
                
                # Adicionar nota sobre disponibilidade
                nota_mapa = doc.add_paragraph()
                nota_mapa.add_run('💡 Nota: ').bold = True
                nota_mapa.add_run('O mapa interativo completo está disponível na versão web do relatório com funcionalidades de zoom, filtros e análise detalhada.')
                
                # Adicionar estatísticas do mapa
                resultados = projeto.get('resultados', {})
                if resultados:
                    stats_mapa_para = doc.add_paragraph()
                    stats_mapa_para.add_run('📊 Estatísticas do Mapa: ').bold = True
                    
                    # CDs no mapa
                    if tipo == 'tradicional':
                        cds_count = len(resultados.get('cds_abertos', []))
                    elif tipo == 'max_cobertura':
                        cds_count = resultados.get('num_cds_selecionados', len(resultados.get('cds_selecionados', [])))
                    else:
                        cds_count = resultados.get('num_cds_selecionados', 0)
                    
                    # Clientes no mapa
                    if tipo == 'tradicional':
                        clientes_unicos = set()
                        for transporte in resultados.get('transportes', []):
                            if transporte.get('destino'):
                                clientes_unicos.add(transporte['destino'])
                        clientes_count = len(clientes_unicos)
                    else:
                        clientes_count = len(resultados.get('atribuicoes', []))
                    
                    stats_mapa_para.add_run(f'{cds_count} CDs e {clientes_count} clientes visualizados no mapa')
                
                # Adicionar linha separadora
                doc.add_paragraph('─' * 50)
            else:
                # Se não houver mapa, adicionar informação
                doc.add_heading('🗺️ Informações de Visualização', level=2)
                sem_mapa_para = doc.add_paragraph()
                sem_mapa_para.add_run('📍 Status: ').bold = True
                sem_mapa_para.add_run('Mapa não gerado para este projeto')
                
                motivo_para = doc.add_paragraph()
                motivo_para.add_run('📋 Motivo: ').bold = True
                motivo_para.add_run('Coordenadas geográficas não disponíveis ou dados insuficientes para geração do mapa')
                
                alternativa_para = doc.add_paragraph()
                alternativa_para.add_run('💡 Alternativa: ').bold = True
                alternativa_para.add_run('Utilize as tabelas de atribuições e análises detalhadas abaixo para visualização dos dados')
            
            # ============= GRID DE ESTATÍSTICAS =============
            doc.add_heading('Métricas Principais', level=2)
            
            cor_tipo = get_tipo_cor(tipo)
            
            # Criar tabela 3x2 para estatísticas
            stats_table = doc.add_table(rows=3, cols=2)
            stats_table.style = 'Medium Grid 1'
            
            # Custo Total / Demanda Coberta
            if tipo == 'tradicional':
                stats_table.cell(0, 0).text = 'Custo Total'
                stats_table.cell(0, 1).text = f'R$ {format_currency_brl(resultados.get("custo_total", 0))}'
            elif tipo == 'max_cobertura':
                stats_table.cell(0, 0).text = 'Demanda Coberta'
                demanda_coberta = resultados.get("demanda_coberta", 0)
                stats_table.cell(0, 1).text = f'{demanda_coberta:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')
            else:
                stats_table.cell(0, 0).text = 'Custo Total'
                custo_ponderado = resultados.get("custo_total_ponderado", 0)
                stats_table.cell(0, 1).text = f'{custo_ponderado:.2f}'
            
            # CDs Ativos
            if tipo == 'tradicional':
                cds_ativos = len(resultados.get('cds_abertos', []))
            elif tipo == 'max_cobertura':
                cds_ativos = resultados.get('num_cds_selecionados', 0)
                if cds_ativos == 0 and resultados.get('cds_selecionados'):
                    cds_ativos = len(resultados['cds_selecionados'])
                elif cds_ativos == 0 and resultados.get('atribuicoes'):
                    cds_unicos = set()
                    for atribuicao in resultados.get('atribuicoes', []):
                        if atribuicao.get('cd_selecionado'):
                            cds_unicos.add(atribuicao.get('cd_selecionado'))
                    cds_ativos = len(cds_unicos)
            else:
                cds_ativos = resultados.get('num_cds_selecionados', 0)
            
            stats_table.cell(1, 0).text = 'CDs Ativos'
            stats_table.cell(1, 1).text = str(cds_ativos)
            
            # Clientes Atendidos
            if tipo == 'tradicional':
                clientes_unicos = set()
                transportes = resultados.get('transportes', [])
                for transporte in transportes:
                    if transporte.get('destino'):
                        clientes_unicos.add(transporte['destino'])
                clientes_atendidos = len(clientes_unicos)
                if clientes_atendidos == 0:
                    clientes_atendidos = len(resultados.get('atribuicoes', []))
            elif tipo == 'max_cobertura':
                clientes_atendidos = len([a for a in resultados.get('atribuicoes', []) if a.get('coberto')])
                if clientes_atendidos == 0:
                    clientes_atendidos = len(resultados.get('atribuicoes', []))
            else:
                clientes_atendidos = len(resultados.get('atribuicoes', []))
            
            stats_table.cell(2, 0).text = 'Clientes Atendidos'
            stats_table.cell(2, 1).text = str(clientes_atendidos)
            
            # Formatar tabela de estatísticas
            for row_idx, row in enumerate(stats_table.rows):
                for col_idx, cell in enumerate(row.cells):
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                            if row_idx == 0:  # Cabeçalho
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                run.font.color.rgb = cor_tipo
            
            # Colorir cabeçalho da tabela (método alternativo)
            try:
                from docx.oxml import parse_xml
                for cell in stats_table.rows[0].cells:
                    # Adicionar cor de fundo ao cabeçalho
                    shading_elm = parse_xml('<w:shd w:fill="2F549D" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
            except:
                pass  # Se falhar, continua sem cor
            
            # ============= SEÇÃO DE CDs SELECIONADOS =============
            cds_lista = None
            if tipo == 'tradicional' and resultados.get('cds_abertos'):
                cds_lista = resultados['cds_abertos']
            elif resultados.get('cds_selecionados'):
                cds_lista = resultados['cds_selecionados']
            
            if cds_lista:
                doc.add_heading('Centros de Distribuição Selecionados', level=2)
                
                # Criar tabela para CDs
                cds_table = doc.add_table(rows=len(cds_lista) + 1, cols=2)
                cds_table.style = 'Medium Grid 1'
                
                # Cabeçalho
                cds_table.cell(0, 0).text = 'Centro de Distribuição'
                cds_table.cell(0, 1).text = 'Informações de Capacidade'
                
                # Formatar cabeçalho
                for cell in cds_table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(11)
                
                # Dados dos CDs
                for idx, cd in enumerate(cds_lista, 1):
                    cds_table.cell(idx, 0).text = str(cd)
                    
                    # Informações de capacidade
                    capacidade_info = "N/A"
                    if tipo == 'tradicional' and resultados.get('capacidade_utilizada') and cd in resultados['capacidade_utilizada']:
                        cap_info = resultados['capacidade_utilizada'][cd]
                        capacidade_info = f"{cap_info.get('utilizada', 0):.0f} / {cap_info.get('disponivel', 0):.0f} ({cap_info.get('percentual_uso', 0):.1f}%)"
                    elif tipo in ['p_medianas', 'p_centros'] and resultados.get('clientes_por_cd') and cd in resultados['clientes_por_cd']:
                        clientes = resultados['clientes_por_cd'][cd]
                        total_clientes = sum(resultados['clientes_por_cd'].values())
                        percentual = (clientes / total_clientes * 100) if total_clientes > 0 else 0
                        capacidade_info = f"{clientes} clientes ({percentual:.1f}%)"
                    
                    cds_table.cell(idx, 1).text = capacidade_info
                    
                    # Formatar células de dados
                    for cell in [cds_table.cell(idx, 0), cds_table.cell(idx, 1)]:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
            
            # ============= ANÁLISE DETALHADA ESPECÍFICA =============
            if tipo == 'tradicional' and resultados:
                doc.add_heading('📊 Análise Detalhada - Modelo Tradicional', level=2)
                
                # Grid de métricas detalhadas
                detalhes_table = doc.add_table(rows=4, cols=2)
                detalhes_table.style = 'Medium Grid 1'
                
                # Custo Total
                detalhes_table.cell(0, 0).text = 'Custo Total'
                detalhes_table.cell(0, 1).text = f'R$ {format_currency_brl(resultados.get("custo_total", 0))}'
                
                # CDs Abertos
                detalhes_table.cell(1, 0).text = 'CDs Abertos'
                detalhes_table.cell(1, 1).text = str(len(resultados.get('cds_abertos', [])))
                
                # Clientes Atendidos
                clientes_unicos = set()
                for transporte in resultados.get('transportes', []):
                    if transporte.get('destino'):
                        clientes_unicos.add(transporte['destino'])
                clientes_atendidos = len(clientes_unicos)
                detalhes_table.cell(2, 0).text = 'Clientes Atendidos'
                detalhes_table.cell(2, 1).text = str(clientes_atendidos)
                
                # Capacidade Usada
                capacidade_total = 0
                if resultados.get('capacidade_utilizada'):
                    for cd, dados in resultados['capacidade_utilizada'].items():
                        capacidade_total += dados.get('utilizada', 0)
                detalhes_table.cell(3, 0).text = 'Capacidade Usada'
                detalhes_table.cell(3, 1).text = f'{capacidade_total:.0f}'
                
                # Formatar tabela detalhada
                for row in detalhes_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(11)
                
                # Análise de Custos (se disponível)
                if resultados.get('custo_fixo_total') or resultados.get('custo_transporte_total'):
                    doc.add_heading('💰 Estrutura Detalhada de Custos', level=3)
                    
                    custos_table = doc.add_table(rows=3, cols=2)
                    custos_table.style = 'Medium Grid 1'
                    
                    # Custos Fixos
                    custos_table.cell(0, 0).text = 'Custos Fixos (CDs)'
                    custos_table.cell(0, 1).text = f'R$ {format_currency_brl(resultados.get("custo_fixo_total", 0))}'
                    
                    # Custos de Transporte
                    custos_table.cell(1, 0).text = 'Custos de Transporte'
                    custos_table.cell(1, 1).text = f'R$ {format_currency_brl(resultados.get("custo_transporte_total", 0))}'
                    
                    # Custo Total
                    custos_table.cell(2, 0).text = 'Custo Total'
                    custos_table.cell(2, 1).text = f'R$ {format_currency_brl(resultados.get("custo_total", 0))}'
                    
                    # Formatar tabela de custos
                    for row in custos_table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
            
            elif tipo == 'p_medianas' and resultados:
                doc.add_heading('📊 Análise Detalhada - Modelo p-Medianas', level=2)
                
                detalhes_pmed_table = doc.add_table(rows=3, cols=2)
                detalhes_pmed_table.style = 'Medium Grid 1'
                
                detalhes_pmed_table.cell(0, 0).text = 'Custo Total Ponderado'
                detalhes_pmed_table.cell(0, 1).text = f'{resultados.get("custo_total_ponderado", 0):.2f}'
                
                detalhes_pmed_table.cell(1, 0).text = 'Número de CDs Selecionados'
                detalhes_pmed_table.cell(1, 1).text = str(resultados.get('num_cds_selecionados', 0))
                
                detalhes_pmed_table.cell(2, 0).text = 'Total de Clientes'
                detalhes_pmed_table.cell(2, 1).text = str(len(resultados.get('atribuicoes', [])))
                
                # Formatar
                for row in detalhes_pmed_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(11)
            
            elif tipo == 'max_cobertura' and resultados:
                doc.add_heading('📊 Análise Detalhada - Máxima Cobertura', level=2)
                
                detalhes_max_table = doc.add_table(rows=4, cols=2)
                detalhes_max_table.style = 'Medium Grid 1'
                
                detalhes_max_table.cell(0, 0).text = 'Demanda Total Coberta'
                detalhes_max_table.cell(0, 1).text = f'{resultados.get("demanda_coberta", 0):.2f}'
                
                detalhes_max_table.cell(1, 0).text = 'Demanda Não Coberta'
                detalhes_max_table.cell(1, 1).text = f'{resultados.get("demanda_nao_coberta", 0):.2f}'
                
                detalhes_max_table.cell(2, 0).text = 'Taxa de Cobertura'
                taxa_cobertura = (resultados.get("demanda_coberta", 0) / (resultados.get("demanda_coberta", 0) + resultados.get("demanda_nao_coberta", 0))) * 100 if (resultados.get("demanda_coberta", 0) + resultados.get("demanda_nao_coberta", 0)) > 0 else 0
                detalhes_max_table.cell(2, 1).text = f'{taxa_cobertura:.1f}%'
                
                detalhes_max_table.cell(3, 0).text = 'Número de CDs'
                detalhes_max_table.cell(3, 1).text = str(cds_ativos)
                
                # Formatar
                for row in detalhes_max_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(11)
            
            elif tipo == 'p_centros' and resultados:
                doc.add_heading('📊 Análise Detalhada - Modelo p-Centros', level=2)
                
                detalhes_pcent_table = doc.add_table(rows=3, cols=2)
                detalhes_pcent_table.style = 'Medium Grid 1'
                
                detalhes_pcent_table.cell(0, 0).text = 'Distância Máxima'
                detalhes_pcent_table.cell(0, 1).text = f'{resultados.get("distancia_maxima", 0):.2f}'
                
                detalhes_pcent_table.cell(1, 0).text = 'Número de CDs Selecionados'
                detalhes_pcent_table.cell(1, 1).text = str(resultados.get('num_cds_selecionados', 0))
                
                detalhes_pcent_table.cell(2, 0).text = 'Total de Clientes'
                detalhes_pcent_table.cell(2, 1).text = str(len(resultados.get('atribuicoes', [])))
                
                # Formatar
                for row in detalhes_pcent_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(11)
            
            # ============= INSIGHT BOXES ESPECÍFICOS =============
            if tipo == 'tradicional' and projeto.get('distribuicao_cds'):
                doc.add_heading('📊 Carga de Trabalho por CD', level=2)
                
                carga_table = doc.add_table(rows=len(projeto['distribuicao_cds']) + 1, cols=3)
                carga_table.style = 'Medium Grid 1'
                
                # Cabeçalho
                carga_table.cell(0, 0).text = 'Centro de Distribuição'
                carga_table.cell(0, 1).text = 'Volume Total'
                carga_table.cell(0, 2).text = '% do Volume'
                
                # Formatar cabeçalho
                for cell in carga_table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                
                # Dados
                for idx, cd_info in enumerate(projeto['distribuicao_cds'], 1):
                    carga_table.cell(idx, 0).text = str(cd_info['nome'])
                    carga_table.cell(idx, 1).text = f"{cd_info['volume']:.0f}"
                    carga_table.cell(idx, 2).text = f"{cd_info['percentual']:.1f}%"
                    
                    # Formatar células
                    for cell in [carga_table.cell(idx, 0), carga_table.cell(idx, 1), carga_table.cell(idx, 2)]:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
            
            elif tipo == 'max_cobertura' and projeto.get('cobertura_stats'):
                doc.add_heading('📊 Análise de Cobertura', level=2)
                
                stats = projeto['cobertura_stats']
                
                cobertura_table = doc.add_table(rows=4, cols=2)
                cobertura_table.style = 'Medium Grid 1'
                
                cobertura_table.cell(0, 0).text = 'Total de Clientes'
                cobertura_table.cell(0, 1).text = str(stats['total_clientes'])
                
                cobertura_table.cell(1, 0).text = 'Clientes Cobertos'
                cobertura_table.cell(1, 1).text = str(stats['clientes_cobertos'])
                
                cobertura_table.cell(2, 0).text = 'Clientes Não Cobertos'
                cobertura_table.cell(2, 1).text = str(stats['clientes_nao_cobertos'])
                
                cobertura_table.cell(3, 0).text = 'Taxa de Cobertura'
                cobertura_table.cell(3, 1).text = f"{stats['taxa_cobertura']:.1f}%"
                
                # Formatar
                for row in cobertura_table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)
                
                # Adicionar raio de cobertura se disponível
                if resultados.get('raio_cobertura'):
                    raio_para = doc.add_paragraph()
                    raio_para.add_run('Raio de Cobertura: ').bold = True
                    raio_para.add_run(f"{resultados['raio_cobertura']} km")
            
            elif tipo in ['p_medianas', 'p_centros'] and resultados.get('clientes_por_cd'):
                doc.add_heading('📊 Distribuição de Clientes por CD', level=2)
                
                clientes_dist = resultados['clientes_por_cd']
                total_clientes = sum(clientes_dist.values())
                
                dist_table = doc.add_table(rows=len(clientes_dist) + 1, cols=3)
                dist_table.style = 'Medium Grid 1'
                
                # Cabeçalho
                dist_table.cell(0, 0).text = 'Centro de Distribuição'
                dist_table.cell(0, 1).text = 'Número de Clientes'
                dist_table.cell(0, 2).text = '% do Total'
                
                # Formatar cabeçalho
                for cell in dist_table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                
                # Dados
                for idx, (cd, clientes) in enumerate(clientes_dist.items(), 1):
                    percentual = (clientes / total_clientes * 100) if total_clientes > 0 else 0
                    dist_table.cell(idx, 0).text = str(cd)
                    dist_table.cell(idx, 1).text = str(clientes)
                    dist_table.cell(idx, 2).text = f"{percentual:.1f}%"
                    
                    # Formatar células
                    for cell in [dist_table.cell(idx, 0), dist_table.cell(idx, 1), dist_table.cell(idx, 2)]:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
            
            # ============= ANÁLISES CRÍTICAS E INSIGHTS =============
            if tipo == 'tradicional' and resultados.get('transportes'):
                # Curva ABC de Custos de Rotas
                doc.add_heading('🔍 Curva ABC dos Custos de Rotas', level=2)
                
                # Calcular custos por rota
                rotas_ordenadas = []
                for transporte in resultados.get('transportes', []):
                    custo_total_rota = transporte.get('quantidade', 0) * transporte.get('custo_unitario', 0)
                    rotas_ordenadas.append({
                        'origem': transporte.get('origem', 'N/A'),
                        'destino': transporte.get('destino', 'N/A'),
                        'custo_total': custo_total_rota,
                        'quantidade': transporte.get('quantidade', 0),
                        'custo_unitario': transporte.get('custo_unitario', 0)
                    })
                
                # Ordenar por custo total
                rotas_ordenadas.sort(key=lambda x: x['custo_total'], reverse=True)
                
                # Top 3 rotas mais críticas
                doc.add_paragraph('📈 Top 3 Rotas Mais Críticas', style='List Bullet')
                for i, rota in enumerate(rotas_ordenadas[:3], 1):
                    rota_para = doc.add_paragraph()
                    rota_para.add_run(f'{i}. {rota["origem"]} → {rota["destino"]}: ').bold = True
                    rota_para.add_run(f'{rota["quantidade"]} unidades × R$ {rota["custo_unitario"]:.2f}/unidade = R$ {rota["custo_total"]:.2f}')
                
                # Análise crítica
                if len(rotas_ordenadas) >= 3 and resultados.get('custo_transporte_total', 0) > 0:
                    custo_top3 = sum(r['custo_total'] for r in rotas_ordenadas[:3])
                    percentual = (custo_top3 / resultados['custo_transporte_total']) * 100
                    critica_para = doc.add_paragraph()
                    critica_para.add_run('⚠️ Análise Crítica: ').bold = True
                    critica_para.add_run(f'As top 3 rotas representam {percentual:.1f}% de todo o custo de transporte da malha.')
            
            # Nível de Importância de cada CD (para todos os tipos)
            if resultados.get('volume_por_cd') or resultados.get('clientes_por_cd'):
                doc.add_heading('📊 Nível de Importância de Cada CD', level=2)
                
                if tipo == 'tradicional':
                    volume_por_cd = resultados.get('volume_por_cd', {})
                    volume_total = sum(volume_por_cd.values())
                    
                    for cd, volume in volume_por_cd.items():
                        percentual = (volume / volume_total * 100) if volume_total > 0 else 0
                        cd_para = doc.add_paragraph()
                        cd_para.add_run(f'CD {cd}: ').bold = True
                        cd_para.add_run(f'{percentual:.1f}% do volume total ({volume:.0f} unidades)')
                
                elif resultados.get('clientes_por_cd'):
                    clientes_por_cd = resultados['clientes_por_cd']
                    total_clientes = sum(clientes_por_cd.values())
                    
                    for cd, clientes in clientes_por_cd.items():
                        percentual = (clientes / total_clientes * 100) if total_clientes > 0 else 0
                        cd_para = doc.add_paragraph()
                        cd_para.add_run(f'CD {cd}: ').bold = True
                        cd_para.add_run(f'{percentual:.1f}% do total ({clientes} clientes)')
            
            # ============= ATRIBUIÇÕES DE CLIENTES =============
            if resultados.get('atribuicoes') and len(resultados['atribuicoes']) > 0:
                doc.add_heading('Atribuições de Clientes', level=2)
                
                # Criar tabela para atribuições (limitar a 20 para não ficar muito grande)
                atribuicoes = resultados['atribuicoes'][:20]  # Limitar para 20 primeiras
                atrib_table = doc.add_table(rows=len(atribuicoes) + 1, cols=3)
                atrib_table.style = 'Medium Grid 1'
                
                # Cabeçalho
                atrib_table.cell(0, 0).text = 'Cliente'
                atrib_table.cell(0, 1).text = 'CD Selecionado'
                atrib_table.cell(0, 2).text = 'Distância'
                
                # Formatar cabeçalho
                for cell in atrib_table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.size = Pt(10)
                
                # Dados das atribuições
                for idx, atribuicao in enumerate(atribuicoes, 1):
                    # Mapear campos diferentes por tipo de modelo
                    if tipo == 'tradicional':
                        # Para modelo tradicional, usar dados dos transportes
                        if resultados.get('transportes') and idx <= len(resultados['transportes']):
                            transporte = resultados['transportes'][idx-1]
                            cliente_nome = transporte.get('destino', f'Cliente {idx}')
                            cd_selecionado = transporte.get('origem', 'N/A')
                            distancia = transporte.get('distancia', 0)
                        else:
                            cliente_nome = f'Cliente {idx}'
                            cd_selecionado = 'N/A'
                            distancia = 0
                    elif tipo == 'max_cobertura':
                        cliente_nome = atribuicao.get('cliente', f'Cliente {idx}')
                        cd_selecionado = atribuicao.get('cd_selecionado', atribuicao.get('cd', f'CD {idx}'))
                        distancia = atribuicao.get('distancia', atribuicao.get('distancia_km', 0))
                    else:  # p_medianas, p_centros
                        # Para p-medianas e p-centros, verificar múltiplos campos possíveis
                        cliente_nome = atribuicao.get('cliente', atribuicao.get('nome_cliente', f'Cliente {idx}'))
                        cd_selecionado = atribuicao.get('cd_selecionado', atribuicao.get('cd', atribuicao.get('cd_origem', f'CD {idx}')))
                        distancia = atribuicao.get('distancia', atribuicao.get('distancia_km', 0))
                    
                    atrib_table.cell(idx, 0).text = str(cliente_nome)
                    atrib_table.cell(idx, 1).text = str(cd_selecionado)
                    atrib_table.cell(idx, 2).text = f'{distancia:.2f}'
                    
                    # Formatar células de dados
                    for cell in [atrib_table.cell(idx, 0), atrib_table.cell(idx, 1), atrib_table.cell(idx, 2)]:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
                
                # Se houver mais de 20 atribuições, adicionar nota
                if len(resultados['atribuicoes']) > 20:
                    nota_para = doc.add_paragraph()
                    nota_para.add_run(f'Nota: Mostrando 20 de {len(resultados["atribuicoes"])} atribuições totais. ')
                    nota_para.add_run('Lista completa disponível na versão web do relatório.').italic = True
        
        # ============= SEÇÃO FINAL - ACHIEVEMENTS =============
        doc.add_page_break()
        doc.add_heading('4. Análise Comparativa e Métricas Consolidadas', level=1)
        
        # Tabela comparativa geral
        doc.add_paragraph('Comparativo Geral dos Projetos', style='List Bullet')
        comp_table = doc.add_table(rows=len(projetos) + 1, cols=5)
        comp_table.style = 'Medium Grid 1'
        
        # Cabeçalho
        comp_table.cell(0, 0).text = 'Projeto'
        comp_table.cell(0, 1).text = 'Modelo'
        comp_table.cell(0, 2).text = 'Data'
        comp_table.cell(0, 3).text = 'CDs'
        comp_table.cell(0, 4).text = 'Clientes'
        
        # Formatar cabeçalho
        for cell in comp_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
        
        # Dados dos projetos
        for i, projeto in enumerate(projetos, 1):
            resultados = projeto.get('resultados', {})
            tipo = projeto.get('tipo_analise')
            
            comp_table.cell(i, 0).text = projeto.get('nome', f'Projeto {i}')
            comp_table.cell(i, 1).text = get_tipo_nome(tipo)
            comp_table.cell(i, 2).text = projeto.get('data_criacao', 'N/A')[:10]
            
            # Número de CDs por tipo
            if tipo == 'tradicional':
                cds_count = len(resultados.get('cds_abertos', []))
            elif tipo == 'max_cobertura':
                cds_count = resultados.get('num_cds_selecionados', len(resultados.get('cds_selecionados', [])))
            else:
                cds_count = resultados.get('num_cds_selecionados', 0)
            comp_table.cell(i, 3).text = str(cds_count)
            
            # Número de clientes por tipo
            if tipo == 'tradicional':
                clientes_unicos = set()
                for transporte in resultados.get('transportes', []):
                    if transporte.get('destino'):
                        clientes_unicos.add(transporte['destino'])
                clientes_count = len(clientes_unicos)
            else:
                clientes_count = len(resultados.get('atribuicoes', []))
            comp_table.cell(i, 4).text = str(clientes_count)
        
        # Métricas consolidadas
        doc.add_paragraph('Métricas Consolidadas do Período', style='List Bullet')
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Medium Grid 1'
        
        metrics_table.cell(0, 0).text = f'{len(projetos)}'
        metrics_table.cell(0, 1).text = 'Projetos Analisados'
        
        metrics_table.cell(1, 0).text = str(consolidados['total_cds'])
        metrics_table.cell(1, 1).text = 'Total de CDs Ativados'
        
        metrics_table.cell(2, 0).text = str(consolidados['total_clientes'])
        metrics_table.cell(2, 1).text = 'Total de Conexões'
        
        metrics_table.cell(3, 0).text = f'R$ {format_currency_brl(consolidados["total_investimento"])}'
        metrics_table.cell(3, 1).text = 'Investimento Total Otimizado'
        
        # Formatar tabela de métricas
        for row in metrics_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
        
        # Distribuição por tipo
        doc.add_paragraph('Distribuição por Tipo de Análise', style='List Bullet')
        dist_table = doc.add_table(rows=4, cols=2)
        dist_table.style = 'Medium Grid 1'
        
        dist_table.cell(0, 0).text = 'radar Máxima Cobertura'
        dist_table.cell(0, 1).text = '1 projeto (25.0% do total)'
        dist_table.cell(1, 0).text = 'center_focus_strong p-Centros'
        dist_table.cell(1, 1).text = '1 projeto (25.0% do total)'
        dist_table.cell(2, 0).text = 'location_city p-Medianas'
        dist_table.cell(2, 1).text = '1 projeto (25.0% do total)'
        dist_table.cell(3, 0).text = 'inventory_2 Tradicional'
        dist_table.cell(3, 1).text = '1 projeto (25.0% do total)'
        
        # Formatar tabela de distribuição
        for row in dist_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        # Insights estratégicos
        doc.add_paragraph('💡 Insights Estratégicos Consolidados', style='List Bullet')
        insights_para = doc.add_paragraph()
        insights_para.add_run('📈 Tendências Identificadas: ').bold = True
        insights_para.add_run('Preferência por modelos de otimização na estratégia logística')
        
        insights_para2 = doc.add_paragraph()
        insights_para2.add_run('🎯 Recomendações Estratégicas: ').bold = True
        insights_para2.add_run('Manter equidade como prioridade operacional')
        
        doc.add_heading('5. Conclusões e Recomendações', level=1)
        
        # Conclusões gerais
        conclusao_geral = doc.add_paragraph()
        conclusao_geral.add_run('A análise integrada dos resultados obtidos através da aplicação dos modelos matemáticos demonstra a eficácia da plataforma NexusNode na otimização de redes logísticas complexas. Os algoritmos implementados conseguiram identificar soluções que equilibram de forma eficiente os múltiplos objetivos estratégicos das operações logísticas modernas.')
        
        # Recomendações por tipo
        doc.add_heading('🎯 Recomendações Específicas', level=2)
        
        for projeto in projetos:
            tipo = projeto.get('tipo_analise')
            resultados = projeto.get('resultados', {})
            
            if tipo == 'max_cobertura':
                doc.add_paragraph('Máxima Cobertura - 📡 Conclusões do Modelo Máxima Cobertura', style='List Bullet')
                mc_para = doc.add_paragraph()
                mc_para.add_run('Alcançada cobertura de ')
                mc_para.add_run(f'{projeto.get("cobertura_stats", {}).get("taxa_cobertura", 0):.1f}%').bold = True
                mc_para.add_run(' da demanda total')
                
                doc.add_paragraph('🎯 Recomendações Específicas', style='List Bullet')
                doc.add_paragraph('Manter raio de cobertura atual - otimização satisfatória', style='List Bullet')
                doc.add_paragraph('Identificar áreas de alta densidade para priorizar novos CDs', style='List Bullet')
                
            elif tipo == 'p_centros':
                doc.add_paragraph('p-Centros - 🎯 Conclusões do Modelo p-Centros', style='List Bullet')
                pc_para = doc.add_paragraph()
                pc_para.add_run('Minimizada distância geográfica máxima entre CDs e clientes, garantindo equidade no atendimento.')
                
                doc.add_paragraph('🎯 Recomendações Específicas', style='List Bullet')
                doc.add_paragraph('Manter foco na equidade de atendimento como diferencial competitivo', style='List Bullet')
                
            elif tipo == 'p_medianas':
                doc.add_paragraph('p-Medianas - 📍 Conclusões do Modelo p-Medianas', style='List Bullet')
                pm_para = doc.add_paragraph()
                pm_para.add_run('Selecionados ')
                pm_para.add_run(f'{resultados.get("num_cds_selecionados", 0)} CDs').bold = True
                pm_para.add_run(' para minimização de custos de transporte')
                
                doc.add_paragraph('🎯 Recomendações Específicas', style='List Bullet')
                doc.add_paragraph('Manter configuração atual dos CDs otimizados', style='List Bullet')
                
            elif tipo == 'tradicional':
                doc.add_paragraph('Otimizador de CDS - 📊 Conclusões do Modelo Tradicional', style='List Bullet')
                trad_para = doc.add_paragraph()
                trad_para.add_run('Foram abertos ')
                trad_para.add_run(f'{len(resultados.get("cds_abertos", []))} centros').bold = True
                trad_para.add_run(' de distribuição otimizando custos fixos')
                
                doc.add_paragraph('🎯 Recomendações Específicas', style='List Bullet')
                doc.add_paragraph('Manter os CDs ativos para continuidade operacional', style='List Bullet')
                doc.add_paragraph('Monitorar custos de transporte para identificar oportunidades de otimização adicional', style='List Bullet')
        
        # ============= SEÇÃO FINAL - PRÓXIMOS PASSOS =============
        doc.add_heading('6. Próximos Passos', level=1)
        
        passos = [
            "Implementação das soluções otimizadas em ambiente produtivo",
            "Monitoramento contínuo de performance e KPIs operacionais", 
            "Ajustes finos nos parâmetros dos modelos conforme feedback",
            "Expansão da análise para novas regiões geográficas",
            "Integração com sistemas ERP e WMS existentes"
        ]
        
        for passo in passos:
            doc.add_paragraph(passo, style='List Bullet')
        
        # Frase final consolidada
        frase_final = doc.add_paragraph()
        frase_final.add_run('Principais Achievements Consolidados: ').bold = True
        frase_final.add_run(f'Análise de {len(projetos)} projetos resultando em {consolidados["total_cds"]} CDs otimizados, atendendo {consolidados["total_clientes"]} pontos de demanda com investimento total otimizado de R$ {format_currency_brl(consolidados["total_investimento"])}.')
        
        # ============= SALVAR DOCUMENTO =============
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            temp_path = tmp_file.name
            doc.save(temp_path)
        
        # Ler o arquivo e enviar
        with open(temp_path, 'rb') as f:
            file_data = f.read()
        
        # Remover arquivo temporário
        os.unlink(temp_path)
        
        # Criar resposta
        output = io.BytesIO(file_data)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'Relatorio_Tecnico_Anual_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        )
        
    except ImportError:
        flash('Biblioteca python-docx não encontrada. Execute: pip install python-docx', 'error')
        return redirect(url_for('dashboard'))
    except Exception as e:
        flash(f'Erro ao gerar relatório Word: {str(e)}', 'error')
        return redirect(url_for('dashboard'))

if __name__ == '__main__':
    print("🚀 Iniciando servidor Flask...")
    print("📁 Pasta atual:", os.getcwd())
    print("🌐 Servidor será iniciado em http://127.0.0.1:5000")
    print("⏳ Aguarde o servidor carregar completamente...")
    
    try:
        app.run(debug=True, host='127.0.0.1', port=5000)
    except Exception as e:
        print(f"❌ Erro ao iniciar servidor: {e}")
        print("🔍 Verificando dependências...")
        
        # Testar imports principais
        try:
            import flask
            print("✅ Flask importado com sucesso")
        except ImportError:
            print("❌ Flask não encontrado - execute: pip install flask")
            
        try:
            import pandas
            print("✅ Pandas importado com sucesso")
        except ImportError:
            print("❌ Pandas não encontrado - execute: pip install pandas")
            
        try:
            import pulp
            print("✅ PuLP importado com sucesso")
        except ImportError:
            print("❌ PuLP não encontrado - execute: pip install pulp")
            
        try:
            import folium
            print("✅ Folium importado com sucesso")
        except ImportError:
            print("❌ Folium não encontrado - execute: pip install folium")

if __name__ == "__main__":
    # O use_reloader=False ajuda a evitar processos duplicados que travam a porta
    app.run(debug=True, port=5000, use_reloader=False)
